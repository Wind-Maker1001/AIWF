from __future__ import annotations

import json
import os
import platform
from typing import Any, Dict, List, Optional, Tuple


def office_lang(params: Optional[Dict[str, Any]] = None) -> str:
    p = params or {}
    lang = str(p.get("office_lang") or "zh").strip().lower()
    if lang not in {"zh", "en"}:
        return "zh"
    return lang


def office_text(zh: str, en: str, params: Optional[Dict[str, Any]] = None) -> str:
    return en if office_lang(params) == "en" else zh


def office_theme_file_path() -> str:
    env_path = os.getenv("AIWF_OFFICE_THEME_FILE")
    if env_path:
        return env_path
    here = os.path.dirname(__file__)
    return os.path.normpath(os.path.join(here, "..", "..", "..", "rules", "templates", "office_themes.json"))


def office_layout_file_path() -> str:
    env_path = os.getenv("AIWF_OFFICE_LAYOUT_FILE")
    if env_path:
        return env_path
    here = os.path.dirname(__file__)
    return os.path.normpath(os.path.join(here, "..", "..", "..", "rules", "templates", "office_layouts.json"))


def load_theme_presets() -> Dict[str, Dict[str, Dict[str, Any]]]:
    defaults = {
        "zh": {
            "professional": {
                "display_name": "专业版",
                "report_title": "AIWF 数据清洗审计报告",
                "cover_title": "AIWF 清洗结果总览",
                "primary_hex": "1F4E78",
                "secondary_hex": "2C6A9D",
                "accent_hex": "F39C34",
                "bg_rgb": [248, 250, 252],
            },
            "academic": {
                "display_name": "学术版",
                "report_title": "学术数据处理报告",
                "cover_title": "研究数据摘要",
                "primary_hex": "2F3E63",
                "secondary_hex": "445A8D",
                "accent_hex": "E3A72F",
                "bg_rgb": [246, 247, 251],
            },
            "debate": {
                "display_name": "辩论版",
                "report_title": "辩论证据简报",
                "cover_title": "辩论证据看板",
                "primary_hex": "7A1F1F",
                "secondary_hex": "9C2F2F",
                "accent_hex": "E9B949",
                "bg_rgb": [252, 247, 247],
            },
            "assignment": {
                "display_name": "作业版",
                "report_title": "课程作业数据报告",
                "cover_title": "作业结果与要点",
                "primary_hex": "245C73",
                "secondary_hex": "2F7C94",
                "accent_hex": "F2A33B",
                "bg_rgb": [245, 251, 253],
            },
            "debate_plus": {
                "display_name": "辩论增强版",
                "report_title": "辩论论据证据报告",
                "cover_title": "辩论证据策略板",
                "primary_hex": "5E1B4F",
                "secondary_hex": "7C2868",
                "accent_hex": "D9A441",
                "bg_rgb": [251, 246, 250],
            },
            "business": {
                "display_name": "商务版",
                "report_title": "业务数据清洗与洞察报告",
                "cover_title": "业务指标简报",
                "primary_hex": "1B4D3E",
                "secondary_hex": "2D7561",
                "accent_hex": "E1A23A",
                "bg_rgb": [244, 250, 247],
            },
        },
        "en": {
            "professional": {
                "display_name": "Professional",
                "report_title": "AIWF Data Cleaning Audit Report",
                "cover_title": "AIWF Cleaning Output Summary",
                "primary_hex": "1F4E78",
                "secondary_hex": "2C6A9D",
                "accent_hex": "F39C34",
                "bg_rgb": [248, 250, 252],
            },
            "academic": {
                "display_name": "Academic",
                "report_title": "Academic Data Processing Report",
                "cover_title": "Research Data Summary",
                "primary_hex": "2F3E63",
                "secondary_hex": "445A8D",
                "accent_hex": "E3A72F",
                "bg_rgb": [246, 247, 251],
            },
            "debate": {
                "display_name": "Debate",
                "report_title": "Debate Evidence Briefing",
                "cover_title": "Debate Evidence Dashboard",
                "primary_hex": "7A1F1F",
                "secondary_hex": "9C2F2F",
                "accent_hex": "E9B949",
                "bg_rgb": [252, 247, 247],
            },
            "assignment": {
                "display_name": "Assignment",
                "report_title": "Course Assignment Data Report",
                "cover_title": "Assignment Results & Highlights",
                "primary_hex": "245C73",
                "secondary_hex": "2F7C94",
                "accent_hex": "F2A33B",
                "bg_rgb": [245, 251, 253],
            },
            "debate_plus": {
                "display_name": "Debate Plus",
                "report_title": "Debate Evidence & Argument Report",
                "cover_title": "Debate Strategy Board",
                "primary_hex": "5E1B4F",
                "secondary_hex": "7C2868",
                "accent_hex": "D9A441",
                "bg_rgb": [251, 246, 250],
            },
            "business": {
                "display_name": "Business",
                "report_title": "Business Data Cleaning & Insight Report",
                "cover_title": "Business KPI Brief",
                "primary_hex": "1B4D3E",
                "secondary_hex": "2D7561",
                "accent_hex": "E1A23A",
                "bg_rgb": [244, 250, 247],
            },
        },
    }
    cfg_path = office_theme_file_path()
    if not os.path.isfile(cfg_path):
        return defaults
    try:
        with open(cfg_path, "r", encoding="utf-8-sig") as f:
            obj = json.load(f)
        if isinstance(obj, dict) and isinstance(obj.get("zh"), dict) and isinstance(obj.get("en"), dict):
            return obj  # type: ignore[return-value]
    except Exception:
        pass
    return defaults


def load_layout_presets() -> Dict[str, Dict[str, Dict[str, Any]]]:
    defaults = {
        "zh": {
            "default": {"docx_max_table_rows": 20, "pptx_max_items": 6},
            "assignment": {"docx_max_table_rows": 24, "pptx_max_items": 8},
        },
        "en": {
            "default": {"docx_max_table_rows": 20, "pptx_max_items": 6},
            "assignment": {"docx_max_table_rows": 24, "pptx_max_items": 8},
        },
    }
    cfg_path = office_layout_file_path()
    if not os.path.isfile(cfg_path):
        return defaults
    try:
        with open(cfg_path, "r", encoding="utf-8-sig") as f:
            obj = json.load(f)
        if isinstance(obj, dict) and isinstance(obj.get("zh"), dict) and isinstance(obj.get("en"), dict):
            return obj  # type: ignore[return-value]
    except Exception:
        pass
    return defaults


def office_font_name(params: Optional[Dict[str, Any]] = None) -> str:
    return "Microsoft YaHei" if office_lang(params) == "zh" else "Calibri"


def _find_font_file(candidates: List[str]) -> Optional[str]:
    for p in candidates:
        if os.path.isfile(p):
            return p
    return None


def pil_font(size: int, params: Optional[Dict[str, Any]] = None, bold: bool = False) -> Any:
    try:
        from PIL import ImageFont  # type: ignore
    except Exception:
        return None
    sys_name = platform.system().lower()
    if office_lang(params) == "zh":
        if os.name == "nt":
            cands = [
                r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
                r"C:\Windows\Fonts\simhei.ttf",
                r"C:\Windows\Fonts\simsun.ttc",
            ]
        elif sys_name == "darwin":
            cands = [
                "/System/Library/Fonts/PingFang.ttc",
                "/System/Library/Fonts/STHeiti Light.ttc",
                "/Library/Fonts/Arial Unicode.ttf",
            ]
        else:
            cands = [
                "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
                "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                "/usr/share/fonts/truetype/arphic/ukai.ttc",
            ]
    else:
        if os.name == "nt":
            cands = [
                r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
                r"C:\Windows\Fonts\calibri.ttf",
            ]
        elif sys_name == "darwin":
            cands = [
                "/System/Library/Fonts/Supplemental/Arial.ttf",
                "/System/Library/Fonts/Supplemental/Calibri.ttf",
            ]
        else:
            cands = [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
            ]
    font_path = _find_font_file(cands)
    if font_path:
        try:
            return ImageFont.truetype(font_path, size=size)
        except Exception:
            return None
    return None


def docx_apply_font(doc: Any, font_name: str) -> None:
    try:
        from docx.oxml.ns import qn  # type: ignore
    except Exception:
        return

    def apply_run_font(run: Any) -> None:
        run.font.name = font_name
        try:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)  # type: ignore
        except Exception:
            pass

    for p in doc.paragraphs:
        for run in p.runs:
            apply_run_font(run)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        apply_run_font(run)


def ppt_apply_font(prs: Any, font_name: str) -> None:
    for slide in prs.slides:
        for shape in slide.shapes:
            if hasattr(shape, "text_frame") and shape.has_text_frame:
                for p in shape.text_frame.paragraphs:
                    for run in p.runs:
                        run.font.name = font_name
            if hasattr(shape, "table") and shape.has_table:
                for row in shape.table.rows:
                    for cell in row.cells:
                        tf = cell.text_frame
                        for p in tf.paragraphs:
                            for run in p.runs:
                                run.font.name = font_name


def _to_int(v: Any) -> Optional[int]:
    try:
        if v is None or v == "":
            return None
        return int(v)
    except Exception:
        return None


def office_max_rows(params: Optional[Dict[str, Any]] = None) -> int:
    p = params or {}
    v = _to_int(p.get("office_max_rows"))
    if v is None:
        return 5000
    return max(100, min(v, 100000))


def office_rows_subset(rows: List[Dict[str, Any]], params: Optional[Dict[str, Any]] = None) -> Tuple[List[Dict[str, Any]], bool]:
    n = office_max_rows(params)
    if len(rows) <= n:
        return rows, False
    return rows[:n], True


def add_picture_fit(slide: Any, image_path: str, x: Any, y: Any, box_w: Any, box_h: Any) -> None:
    try:
        from PIL import Image  # type: ignore
    except Exception:
        slide.shapes.add_picture(image_path, x, y, width=box_w)
        return

    with Image.open(image_path) as img:
        iw, ih = img.size
    if iw <= 0 or ih <= 0:
        slide.shapes.add_picture(image_path, x, y, width=box_w)
        return

    wr = float(box_w) / float(iw)
    hr = float(box_h) / float(ih)
    scale = min(wr, hr)
    tw = int(iw * scale)
    th = int(ih * scale)
    left = int(x + (box_w - tw) / 2)
    top = int(y + (box_h - th) / 2)
    slide.shapes.add_picture(image_path, left, top, width=tw, height=th)


def office_theme_settings(params: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    p = params or {}
    lang = office_lang(p)
    theme = str(p.get("office_theme") or "professional").strip().lower()
    presets = load_theme_presets()
    lang_map = presets.get(lang) if isinstance(presets.get(lang), dict) else {}
    prof = lang_map.get("professional") if isinstance(lang_map.get("professional"), dict) else {}
    cfg = (lang_map.get(theme) if isinstance(lang_map.get(theme), dict) else prof).copy()
    if not cfg:
        cfg = {
            "display_name": ("专业版" if lang == "zh" else "Professional"),
            "report_title": ("AIWF 数据清洗审计报告" if lang == "zh" else "AIWF Data Cleaning Audit Report"),
            "cover_title": ("AIWF 清洗结果总览" if lang == "zh" else "AIWF Cleaning Output Summary"),
            "primary_hex": "1F4E78",
            "secondary_hex": "2C6A9D",
            "accent_hex": "F39C34",
            "bg_rgb": [248, 250, 252],
        }
    cfg["name"] = theme
    if isinstance(cfg.get("bg_rgb"), list):
        cfg["bg_rgb"] = tuple(cfg["bg_rgb"])
    return cfg


def office_layout_settings(params: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    p = params or {}
    lang = office_lang(p)
    theme = str(p.get("office_theme") or "professional").strip().lower()
    presets = load_layout_presets()
    lang_map = presets.get(lang) if isinstance(presets.get(lang), dict) else {}
    base = lang_map.get("default") if isinstance(lang_map.get("default"), dict) else {}
    ext = lang_map.get(theme) if isinstance(lang_map.get(theme), dict) else {}
    out: Dict[str, Any] = {}
    out.update(base or {})
    out.update(ext or {})
    if not out:
        out = {"docx_max_table_rows": 20, "pptx_max_items": 6}
    return out


def office_quality_mode(params: Optional[Dict[str, Any]] = None) -> str:
    p = params or {}
    mode = str(p.get("office_quality_mode") or "high").strip().lower()
    if mode not in {"high", "standard"}:
        return "high"
    return mode


def office_is_high_quality(params: Optional[Dict[str, Any]] = None) -> bool:
    return office_quality_mode(params) == "high"


def hex_to_rgb(h: str) -> Tuple[int, int, int]:
    s = str(h or "").strip().lstrip("#")
    if len(s) == 3:
        s = "".join([c + c for c in s])
    if len(s) != 6:
        return (31, 78, 120)
    try:
        return (int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))
    except Exception:
        return (31, 78, 120)
