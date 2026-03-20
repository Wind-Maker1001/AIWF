from __future__ import annotations

import os
from typing import Any, Dict, Optional

from aiwf.office_style import (
    docx_apply_font,
    hex_to_rgb,
    office_font_name,
    office_is_high_quality,
    office_layout_settings,
    office_quality_mode,
    office_text,
    office_theme_settings,
)


def write_audit_docx(
    docx_path: str,
    job_id: str,
    profile: Dict[str, Any],
    image_path: Optional[str] = None,
    params: Optional[Dict[str, Any]] = None,
    *,
    utc_now_str,
) -> None:
    from docx import Document  # type: ignore
    from docx.shared import Inches, Pt, RGBColor  # type: ignore

    theme = office_theme_settings(params)
    layout = office_layout_settings(params)
    high_quality = office_is_high_quality(params)
    primary_rgb = hex_to_rgb(str(theme.get("primary_hex", "1F4E78")))
    font_name = office_font_name(params)
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    title = doc.add_heading(str(theme.get("report_title")), level=1)
    title.runs[0].font.size = Pt(20)
    for run in title.runs:
        run.font.name = font_name
        run.font.color.rgb = RGBColor(primary_rgb[0], primary_rgb[1], primary_rgb[2])

    subtitle = doc.add_paragraph(
        office_text(
            "本报告用于作业与辩论场景，展示清洗结果、质量指标和可复用结论。",
            "This report summarizes cleaned output, quality metrics, and reusable conclusions.",
            params,
        )
    )
    subtitle.runs[0].font.name = font_name
    subtitle.runs[0].font.size = Pt(10.5)

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Light List Accent 1"
    meta.cell(0, 0).text = office_text("任务ID", "Job ID", params)
    meta.cell(0, 1).text = job_id
    meta.cell(1, 0).text = office_text("流程步骤", "Step", params)
    meta.cell(1, 1).text = "cleaning"
    meta.cell(2, 0).text = office_text("生成时间", "Generated At", params)
    meta.cell(2, 1).text = utc_now_str()
    meta.cell(3, 0).text = office_text("状态", "Status", params)
    meta.cell(3, 1).text = office_text("完成", "DONE", params)
    meta.add_row().cells[0].text = office_text("主题", "Theme", params)
    meta.rows[-1].cells[1].text = str(theme.get("display_name", theme.get("name")))
    meta.add_row().cells[0].text = office_text("质量模式", "Quality Mode", params)
    meta.rows[-1].cells[1].text = office_quality_mode(params)

    doc.add_paragraph("")
    doc.add_heading(office_text("核心指标", "Core Metrics", params), level=2)
    metrics = doc.add_table(rows=1, cols=2)
    metrics.style = "Light Grid Accent 1"
    metrics.rows[0].cells[0].text = office_text("指标", "Metric", params)
    metrics.rows[0].cells[1].text = office_text("数值", "Value", params)
    core_items = [
        (office_text("行数", "Rows", params), profile.get("rows")),
        (office_text("列数", "Columns", params), profile.get("cols")),
        (office_text("金额总计", "Sum Amount", params), profile.get("sum_amount")),
        (office_text("金额最小值", "Min Amount", params), profile.get("min_amount")),
        (office_text("金额最大值", "Max Amount", params), profile.get("max_amount")),
        (office_text("金额均值", "Avg Amount", params), profile.get("avg_amount")),
    ]
    max_core = max(4, int(layout.get("docx_max_table_rows", 20)))
    for key, value in core_items[:max_core]:
        cells = metrics.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(value)

    quality = profile.get("quality") if isinstance(profile.get("quality"), dict) else {}
    doc.add_heading(office_text("质量摘要", "Quality Summary", params), level=2)
    q_items = [
        (office_text("输入行数", "Input Rows", params), quality.get("input_rows")),
        (office_text("输出行数", "Output Rows", params), quality.get("output_rows")),
        (office_text("无效行数", "Invalid Rows", params), quality.get("invalid_rows")),
        (office_text("过滤行数", "Filtered Rows", params), quality.get("filtered_rows")),
        (office_text("去重移除行数", "Duplicates Removed", params), quality.get("duplicate_rows_removed")),
    ]
    q = doc.add_table(rows=1, cols=2)
    q.style = "Light Grid Accent 1"
    q.rows[0].cells[0].text = office_text("质量指标", "Quality Metric", params)
    q.rows[0].cells[1].text = office_text("数值", "Value", params)
    for key, value in q_items:
        cells = q.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(value)

    doc.add_heading(office_text("结论要点", "Key Takeaways", params), level=2)
    invalid_rows = int(quality.get("invalid_rows", 0) or 0)
    filtered_rows = int(quality.get("filtered_rows", 0) or 0)
    rows_count = int(profile.get("rows", 0) or 0)
    takeaways = [
        office_text(
            f"最终可用数据 {rows_count} 行；可直接用于作业正文与辩论证据表。",
            f"Final usable rows: {rows_count}; ready for assignment writing and debate evidence.",
            params,
        ),
        office_text(
            f"无效行 {invalid_rows}，过滤行 {filtered_rows}；建议在附录中说明处理规则。",
            f"Invalid rows: {invalid_rows}, filtered rows: {filtered_rows}; explain data rules in appendix.",
            params,
        ),
        office_text(
            "图表与表格已按统一主题输出，可直接复制到课程汇报材料。",
            "Charts and tables are exported with a unified theme and can be reused directly.",
            params,
        ),
    ]
    for takeaway in takeaways:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(takeaway)
        run.font.name = font_name

    numeric_stats = profile.get("numeric_stats") if isinstance(profile.get("numeric_stats"), dict) else {}
    if numeric_stats:
        doc.add_heading(office_text("字段统计", "Field Statistics", params), level=2)
        stats_table = doc.add_table(rows=1, cols=5)
        stats_table.style = "Light Grid Accent 1"
        headers = [
            office_text("字段", "Field", params),
            office_text("总和", "Sum", params),
            office_text("最小值", "Min", params),
            office_text("最大值", "Max", params),
            office_text("均值", "Avg", params),
        ]
        for i, header in enumerate(headers):
            stats_table.rows[0].cells[i].text = header
        for field in sorted(numeric_stats.keys())[:12]:
            stats = numeric_stats.get(field) if isinstance(numeric_stats.get(field), dict) else {}
            cells = stats_table.add_row().cells
            cells[0].text = str(field)
            cells[1].text = str(stats.get("sum"))
            cells[2].text = str(stats.get("min"))
            cells[3].text = str(stats.get("max"))
            cells[4].text = str(stats.get("avg"))
    if image_path and os.path.isfile(image_path):
        doc.add_heading(office_text("可视化说明", "Visual Explanation", params), level=2)
        doc.add_picture(image_path, width=Inches(6.2 if high_quality else 6.6))
    docx_apply_font(doc, font_name)
    doc.save(docx_path)
