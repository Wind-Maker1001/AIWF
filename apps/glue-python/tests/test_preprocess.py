import json
import os
import tempfile
import unittest
from unittest.mock import Mock, patch

from aiwf import preprocess


class PreprocessTests(unittest.TestCase):
    def test_validate_preprocess_spec(self):
        ok = preprocess.validate_preprocess_spec(
            {
                "header_map": {"Amt": "amount"},
                "header_mapping_mode": "auto",
                "amount_fields": ["amount"],
                "date_fields": ["biz_date"],
                "amount_round_digits": 2,
            }
        )
        self.assertTrue(ok["ok"])

        bad = preprocess.validate_preprocess_spec({"amount_round_digits": 100, "amount_fields": "x"})
        self.assertFalse(bad["ok"])
        warn = preprocess.validate_preprocess_spec({"unknown_key": 1})
        self.assertTrue(warn["ok"])
        self.assertTrue(any("unknown preprocess keys" in w for w in warn["warnings"]))
        bad2 = preprocess.validate_preprocess_spec({"on_file_error": "x"})
        self.assertFalse(bad2["ok"])
        bad3 = preprocess.validate_preprocess_spec({"generate_quality_report": "yes"})
        self.assertFalse(bad3["ok"])
        bad4 = preprocess.validate_preprocess_spec({"chunk_mode": "abc"})
        self.assertFalse(bad4["ok"])
        bad5 = preprocess.validate_preprocess_spec({"ocr_lang": 1})
        self.assertFalse(bad5["ok"])
        bad6 = preprocess.validate_preprocess_spec({"ocr_config": 1})
        self.assertFalse(bad6["ok"])
        bad7 = preprocess.validate_preprocess_spec({"ocr_preprocess": 1})
        self.assertFalse(bad7["ok"])
        bad8 = preprocess.validate_preprocess_spec({"ocr_preprocess": "bad"})
        self.assertFalse(bad8["ok"])
        bad9 = preprocess.validate_preprocess_spec({"pipeline": {"stages": []}})
        self.assertFalse(bad9["ok"])
        bad10 = preprocess.validate_preprocess_spec({"export_canonical_bundle": "yes"})
        self.assertFalse(bad10["ok"])
        bad11 = preprocess.validate_preprocess_spec({"field_transforms": [{"field": "text", "op": "missing_op"}]})
        self.assertFalse(bad11["ok"])
        bad12 = preprocess.validate_preprocess_spec({"row_filters": [{"field": "text", "op": "missing_filter"}]})
        self.assertFalse(bad12["ok"])
        bad13 = preprocess.validate_preprocess_spec({"header_mapping_mode": "invalid"})
        self.assertFalse(bad13["ok"])
        ok2 = preprocess.validate_preprocess_spec(
            {
                "external_enrichment_mode": "auto",
                "document_parse_backend": "local",
                "citation_parse_backend": "regex",
                "url_metadata_enrichment": True,
                "pdf_text_fast_path": True,
                "pdf_text_fast_path_min_rows": 1,
                "pdf_text_fast_path_min_chars": 40,
            }
        )
        self.assertTrue(ok2["ok"])
        bad_pdf_fast_path = preprocess.validate_preprocess_spec(
            {"pdf_text_fast_path": "yes", "pdf_text_fast_path_min_chars": 0}
        )
        self.assertFalse(bad_pdf_fast_path["ok"])
        self.assertIn("pdf_text_fast_path must be boolean", bad_pdf_fast_path["errors"])
        self.assertIn("pdf_text_fast_path_min_chars must be > 0", bad_pdf_fast_path["errors"])
        warn2 = preprocess.validate_preprocess_spec(
            {
                "standardize_evidence": True,
                "deduplicate_by": ["source_path", "text", "url"],
            }
        )
        self.assertTrue(warn2["ok"])
        self.assertTrue(any("text->claim_text" in w and "url->source_url" in w for w in warn2["warnings"]))

    def test_preprocess_passes_ocr_options_to_ingest(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "a.txt")
            out = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("placeholder")

            with patch("aiwf.preprocess.ingest.load_rows_from_files") as load_rows:
                load_rows.return_value = ([{"text": "x"}], {"input_format": "txt", "skipped_files": [], "failed_files": []})
                preprocess.preprocess_file(
                    src,
                    out,
                    {
                        "input_files": [src],
                        "output_format": "jsonl",
                        "ocr_enabled": True,
                        "ocr_lang": "chi_sim+eng",
                        "ocr_config": "--oem 1 --psm 6",
                        "ocr_preprocess": "adaptive",
                    },
                )

                _, kwargs = load_rows.call_args
                self.assertEqual(kwargs["ocr_lang"], "chi_sim+eng")
                self.assertEqual(kwargs["ocr_config"], "--oem 1 --psm 6")
                self.assertEqual(kwargs["ocr_preprocess"], "adaptive")

    def test_preprocess_csv_file(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.csv")
            dst = os.path.join(tmp, "cooked.csv")
            with open(src, "w", encoding="utf-8", newline="\n") as f:
                f.write("ID,Amt,Date\n")
                f.write("1,$1,200.56,2026/02/16\n")
                f.write("2,  99.1 ,2026-02-17\n")

            # fix comma issue in amount field by quoting in source
            with open(src, "w", encoding="utf-8", newline="\n") as f:
                f.write('ID,Amt,Date\n')
                f.write('1,"$1,200.56",2026/02/16\n')
                f.write('2,"  99.1 ",2026-02-17\n')

            res = preprocess.preprocess_csv_file(
                src,
                dst,
                {
                    "header_map": {"ID": "id", "Amt": "amount", "Date": "biz_date"},
                    "amount_fields": ["amount"],
                    "date_fields": ["biz_date"],
                    "date_input_formats": ["%Y/%m/%d", "%Y-%m-%d"],
                },
            )
            self.assertTrue(os.path.isfile(dst))
            self.assertEqual(res["summary"]["input_rows"], 2)
            self.assertEqual(res["summary"]["output_rows"], 2)

            rows, _ = preprocess._read_csv(dst)
            self.assertEqual(rows[0]["id"], "1")
            self.assertEqual(rows[0]["amount"], "1200.56")
            self.assertEqual(rows[0]["biz_date"], "2026-02-16")

    def test_preprocess_csv_file_for_bank_statement_profile(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "bank.csv")
            dst = os.path.join(tmp, "bank_out.jsonl")
            with open(src, "w", encoding="utf-8", newline="\n") as f:
                f.write("账号,交易日期,借方金额,贷方金额,余额,对方户名,流水号\n")
                f.write('62220001,2026/03/01,120.5,0,"1,250.00",张三,TXN-001\n')

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_format": "csv",
                    "output_format": "jsonl",
                    "canonical_profile": "bank_statement",
                    "header_map": {
                        "账号": "account_no",
                        "交易日期": "txn_date",
                        "借方金额": "debit_amount",
                        "贷方金额": "credit_amount",
                        "余额": "balance",
                        "对方户名": "counterparty_name",
                        "流水号": "ref_no",
                    },
                    "amount_fields": ["debit_amount", "credit_amount", "balance"],
                    "date_fields": ["txn_date"],
                },
            )
            self.assertEqual(res["summary"]["output_rows"], 1)
            rows = preprocess._read_jsonl(dst)
            self.assertEqual(rows[0]["account_no"], "62220001")
            self.assertEqual(rows[0]["txn_date"], "2026-03-01")
            self.assertEqual(rows[0]["debit_amount"], 120.5)
            self.assertEqual(rows[0]["balance"], 1250.0)

    def test_preprocess_jsonl_with_transforms_and_filters(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "cooked.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(json.dumps({"speaker": " Alice ", "text": "Visit https://x.com NOW", "score": "9.6"}) + "\n")
                f.write(json.dumps({"speaker": "bob", "text": "contact me a@b.com", "score": "3.1"}) + "\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_format": "jsonl",
                    "output_format": "jsonl",
                    "field_transforms": [
                        {"field": "speaker", "op": "trim"},
                        {"field": "speaker", "op": "lower"},
                        {"field": "text", "op": "remove_urls"},
                        {"field": "text", "op": "remove_emails"},
                        {"field": "text", "op": "collapse_whitespace"},
                        {"field": "score", "op": "parse_number"},
                        {"field": "score", "op": "round_number", "digits": 0},
                    ],
                    "row_filters": [{"field": "score", "op": "gte", "value": 5}],
                },
            )
            self.assertTrue(os.path.isfile(dst))
            self.assertEqual(res["output_format"], "jsonl")
            self.assertEqual(res["summary"]["output_rows"], 1)
            rows = preprocess._read_jsonl(dst)
            self.assertEqual(rows[0]["speaker"], "alice")
            self.assertNotIn("https://", rows[0]["text"])
            self.assertEqual(rows[0]["score"], 10.0)

    def test_preprocess_can_use_rust_v2_for_supported_transforms(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "cooked.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(json.dumps({"Speaker": " Alice ", "biz_date": "2026/03/01"}) + "\n")

            transform_resp = Mock()
            transform_resp.status_code = 200
            transform_resp.json.return_value = {
                "ok": True,
                "operator": "transform_rows_v3",
                "rows": [{"speaker": "alice", "biz_date": "2026-03-01"}],
                "quality": {
                    "input_rows": 1,
                    "output_rows": 1,
                    "invalid_rows": 0,
                    "filtered_rows": 0,
                    "duplicate_rows_removed": 0,
                    "numeric_cells_total": 0,
                    "numeric_cells_parsed": 0,
                    "date_cells_total": 1,
                    "date_cells_parsed": 1,
                },
                "trace_id": "tp1",
                "audit": {"schema": "transform_rows_v2.audit.v1"},
            }
            quality_resp = Mock()
            quality_resp.status_code = 200
            quality_resp.json.return_value = {
                "ok": True,
                "passed": True,
                "report": {"violations": [], "metrics": {}},
            }
            with patch("requests.post", side_effect=[transform_resp, quality_resp]):
                res = preprocess.preprocess_file(
                    src,
                    dst,
                    {
                        "input_format": "jsonl",
                        "output_format": "jsonl",
                        "use_rust_v2": True,
                        "field_transforms": [
                            {"field": "speaker", "op": "trim"},
                            {"field": "speaker", "op": "lower"},
                            {"field": "biz_date", "op": "parse_date"},
                        ],
                    },
                )
            rows = preprocess._read_jsonl(dst)
            self.assertEqual(rows[0]["speaker"], "alice")
            self.assertEqual(rows[0]["biz_date"], "2026-03-01")
            self.assertEqual(res["summary"]["cleaning_spec_version"], "cleaning_spec.v2")
            self.assertTrue(res["summary"]["rust_v2_used"])
            self.assertEqual(res["execution_mode"], "rust_v3")
            self.assertEqual(res["eligibility_reason"], "eligible")
            self.assertEqual(res["execution_audit"]["schema"], "transform_rows_v2.audit.v1")
            self.assertEqual(res["summary"]["execution_mode"], "rust_v3")
            self.assertEqual(res["summary"]["row_transform_engine"], "transform_rows_v3")
            self.assertEqual(res["execution_audit"]["stage_plan"]["schema_version"], "preprocess_stage_plan.v1")
            self.assertEqual(
                [item["name"] for item in res["execution_audit"]["stage_plan"]["stages"]],
                ["row_transform", "standardize_evidence", "chunk_text", "detect_conflicts", "quality_check", "materialize"],
            )

    def test_preprocess_rust_v2_blocked_by_chunk_mode(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "cooked.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(json.dumps({"text": "a. b."}) + "\n")

            transform_resp = Mock()
            transform_resp.status_code = 200
            transform_resp.json.return_value = {
                "ok": True,
                "operator": "transform_rows_v3",
                "rows": [{"text": "a. b."}],
                "quality": {"input_rows": 1, "output_rows": 1, "invalid_rows": 0, "filtered_rows": 0, "duplicate_rows_removed": 0},
                "trace_id": "tp-chunk",
                "audit": {"schema": "transform_rows_v2.audit.v1"},
            }
            postprocess_resp = Mock()
            postprocess_resp.status_code = 200
            postprocess_resp.json.return_value = {
                "ok": True,
                "operator": "postprocess_rows_v1",
                "rows": [{"text": "a.", "chunk_seq": 0}, {"text": "b.", "chunk_seq": 1}],
                "quality": {"input_rows": 1, "output_rows": 2, "standardized_rows": 0, "chunked_rows_created": 1, "conflict_rows_marked": 0},
                "trace_id": "pp-chunk",
                "audit": {"schema": "postprocess_rows_v1.audit.v1"},
            }
            quality_resp = Mock()
            quality_resp.status_code = 200
            quality_resp.json.return_value = {"ok": True, "passed": True, "report": {"violations": [], "metrics": {}}}

            with patch("requests.post", side_effect=[transform_resp, postprocess_resp, quality_resp]) as post:
                res = preprocess.preprocess_file(
                    src,
                    dst,
                    {
                        "input_format": "jsonl",
                        "output_format": "jsonl",
                        "use_rust_v2": True,
                        "chunk_mode": "sentence",
                    },
                )
            self.assertEqual(post.call_count, 3)
            self.assertEqual(res["execution_mode"], "rust_v3_postprocess_v1")
            self.assertEqual(res["eligibility_reason"], "mixed_rust_postprocess")
            self.assertEqual(res["summary"]["execution_audit"]["schema"], "transform_rows_v2.audit.v1")
            self.assertEqual(res["summary"]["postprocess_engine"], "postprocess_rows_v1")
            self.assertEqual(res["summary"]["chunked_rows_created"], 1)
            self.assertTrue(res["execution_audit"]["stage_plan"]["stages"][2]["enabled"])

    def test_preprocess_rust_v2_blocked_by_conflict_detection(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "cooked.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(json.dumps({"claim_text": "support tax"}) + "\n")

            transform_resp = Mock()
            transform_resp.status_code = 200
            transform_resp.json.return_value = {
                "ok": True,
                "operator": "transform_rows_v3",
                "rows": [{"claim_text": "support tax"}],
                "quality": {"input_rows": 1, "output_rows": 1, "invalid_rows": 0, "filtered_rows": 0, "duplicate_rows_removed": 0},
                "trace_id": "tp-conflict",
                "audit": {"schema": "transform_rows_v2.audit.v1"},
            }
            postprocess_resp = Mock()
            postprocess_resp.status_code = 200
            postprocess_resp.json.return_value = {
                "ok": True,
                "operator": "postprocess_rows_v1",
                "rows": [{"claim_text": "support tax", "conflict_flag": False, "conflict_topic": "tax", "conflict_polarity": "pro"}],
                "quality": {"input_rows": 1, "output_rows": 1, "standardized_rows": 0, "chunked_rows_created": 0, "conflict_rows_marked": 0},
                "trace_id": "pp-conflict",
                "audit": {"schema": "postprocess_rows_v1.audit.v1"},
            }
            quality_resp = Mock()
            quality_resp.status_code = 200
            quality_resp.json.return_value = {"ok": True, "passed": True, "report": {"violations": [], "metrics": {}}}

            with patch("requests.post", side_effect=[transform_resp, postprocess_resp, quality_resp]) as post:
                res = preprocess.preprocess_file(
                    src,
                    dst,
                    {
                        "input_format": "jsonl",
                        "output_format": "jsonl",
                        "use_rust_v2": True,
                        "detect_conflicts": True,
                    },
                )
            self.assertEqual(post.call_count, 3)
            self.assertEqual(res["execution_mode"], "rust_v3_postprocess_v1")
            self.assertEqual(res["eligibility_reason"], "mixed_rust_postprocess")
            self.assertEqual(res["summary"]["conflict_rows_marked"], 0)
            self.assertTrue(res["execution_audit"]["stage_plan"]["stages"][3]["enabled"])

    def test_preprocess_rust_v2_blocked_by_standardize_evidence(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "cooked.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(json.dumps({"text": "claim"}) + "\n")

            transform_resp = Mock()
            transform_resp.status_code = 200
            transform_resp.json.return_value = {
                "ok": True,
                "operator": "transform_rows_v3",
                "rows": [{"text": "claim"}],
                "quality": {"input_rows": 1, "output_rows": 1, "invalid_rows": 0, "filtered_rows": 0, "duplicate_rows_removed": 0},
                "trace_id": "tp-std",
                "audit": {"schema": "transform_rows_v2.audit.v1"},
            }
            postprocess_resp = Mock()
            postprocess_resp.status_code = 200
            postprocess_resp.json.return_value = {
                "ok": True,
                "operator": "postprocess_rows_v1",
                "rows": [{"evidence_id": "e1", "claim_text": "claim", "source_path": src, "source_type": "jsonl"}],
                "quality": {"input_rows": 1, "output_rows": 1, "standardized_rows": 1, "chunked_rows_created": 0, "conflict_rows_marked": 0},
                "trace_id": "pp-std",
                "audit": {"schema": "postprocess_rows_v1.audit.v1"},
            }
            quality_resp = Mock()
            quality_resp.status_code = 200
            quality_resp.json.return_value = {"ok": True, "passed": True, "report": {"violations": [], "metrics": {}}}

            with patch("requests.post", side_effect=[transform_resp, postprocess_resp, quality_resp]) as post:
                res = preprocess.preprocess_file(
                    src,
                    dst,
                    {
                        "input_format": "jsonl",
                        "output_format": "jsonl",
                        "use_rust_v2": True,
                        "standardize_evidence": True,
                    },
                )
            self.assertEqual(post.call_count, 3)
            self.assertEqual(res["execution_mode"], "rust_v3_postprocess_v1")
            self.assertEqual(res["eligibility_reason"], "mixed_rust_postprocess")
            self.assertEqual(res["summary"]["standardized_rows"], 1)
            self.assertTrue(res["execution_audit"]["stage_plan"]["stages"][1]["enabled"])

    def test_preprocess_rust_v2_blocked_by_custom_transform(self):
        def prefix_transform(value, cfg):
            if value is None:
                return value, False
            return f"{cfg.get('prefix', '')}{value}", True

        preprocess.register_field_transform(
            "prefix_blocker",
            prefix_transform,
            domain="custom-preprocess",
            domain_metadata={"label": "Custom Preprocess", "backend": "extension", "builtin": False},
        )
        try:
            with tempfile.TemporaryDirectory() as tmp:
                src = os.path.join(tmp, "raw.jsonl")
                dst = os.path.join(tmp, "cooked.jsonl")
                with open(src, "w", encoding="utf-8") as f:
                    f.write(json.dumps({"speaker": "alice"}) + "\n")

                with patch("requests.post") as post:
                    res = preprocess.preprocess_file(
                        src,
                        dst,
                        {
                            "input_format": "jsonl",
                            "output_format": "jsonl",
                            "use_rust_v2": True,
                            "field_transforms": [{"field": "speaker", "op": "prefix_blocker", "prefix": "team-"}],
                        },
                    )
                post.assert_not_called()
                self.assertEqual(res["execution_mode"], "python_legacy")
                self.assertEqual(res["eligibility_reason"], "unsupported_field_transform")
        finally:
            preprocess.unregister_field_transform("prefix_blocker")

    def test_preprocess_rust_v2_blocked_by_custom_filter(self):
        def starts_with_filter(row, cfg):
            return str(row.get(str(cfg.get("field") or "")) or "").startswith(str(cfg.get("value") or ""))

        preprocess.register_row_filter(
            "starts_with_blocker",
            starts_with_filter,
            domain="custom-preprocess",
            domain_metadata={"label": "Custom Preprocess", "backend": "extension", "builtin": False},
        )
        try:
            with tempfile.TemporaryDirectory() as tmp:
                src = os.path.join(tmp, "raw.jsonl")
                dst = os.path.join(tmp, "cooked.jsonl")
                with open(src, "w", encoding="utf-8") as f:
                    f.write(json.dumps({"speaker": "alice"}) + "\n")

                with patch("requests.post") as post:
                    res = preprocess.preprocess_file(
                        src,
                        dst,
                        {
                            "input_format": "jsonl",
                            "output_format": "jsonl",
                            "use_rust_v2": True,
                            "row_filters": [{"field": "speaker", "op": "starts_with_blocker", "value": "a"}],
                        },
                    )
                post.assert_not_called()
                self.assertEqual(res["execution_mode"], "python_legacy")
                self.assertEqual(res["eligibility_reason"], "unsupported_row_filter")
        finally:
            preprocess.unregister_row_filter("starts_with_blocker")

    def test_preprocess_supports_registered_custom_transform_and_filter(self):
        def prefix_transform(value, cfg):
            if value is None:
                return value, False
            return f"{cfg.get('prefix', '')}{value}", True

        def starts_with_filter(row, cfg):
            field = str(cfg.get("field") or "").strip()
            if not field:
                return True
            return str(row.get(field) or "").startswith(str(cfg.get("value") or ""))

        preprocess.register_field_transform(
            "prefix",
            prefix_transform,
            domain="custom-preprocess",
            domain_metadata={"label": "Custom Preprocess", "backend": "extension", "builtin": False},
        )
        preprocess.register_row_filter(
            "starts_with",
            starts_with_filter,
            domain="custom-preprocess",
            domain_metadata={"label": "Custom Preprocess", "backend": "extension", "builtin": False},
        )
        try:
            with tempfile.TemporaryDirectory() as tmp:
                src = os.path.join(tmp, "raw.jsonl")
                dst = os.path.join(tmp, "cooked.jsonl")
                with open(src, "w", encoding="utf-8") as f:
                    f.write(json.dumps({"speaker": "alice"}) + "\n")
                    f.write(json.dumps({"speaker": "bob"}) + "\n")

                valid = preprocess.validate_preprocess_spec(
                    {
                        "input_format": "jsonl",
                        "output_format": "jsonl",
                        "field_transforms": [{"field": "speaker", "op": "prefix", "prefix": "team-"}],
                        "row_filters": [{"field": "speaker", "op": "starts_with", "value": "team-a"}],
                    }
                )
                self.assertTrue(valid["ok"])

                res = preprocess.preprocess_file(
                    src,
                    dst,
                    {
                        "input_format": "jsonl",
                        "output_format": "jsonl",
                        "field_transforms": [{"field": "speaker", "op": "prefix", "prefix": "team-"}],
                        "row_filters": [{"field": "speaker", "op": "starts_with", "value": "team-a"}],
                    },
                )
                rows = preprocess._read_jsonl(dst)
                field_details = {item["op"]: item for item in preprocess.list_field_transform_details()}
                row_domains = preprocess.list_row_filter_domains()
        finally:
            preprocess.unregister_field_transform("prefix")
            preprocess.unregister_row_filter("starts_with")

        self.assertEqual(res["summary"]["output_rows"], 1)
        self.assertEqual(rows[0]["speaker"], "team-alice")
        self.assertEqual(field_details["prefix"]["domain"], "custom-preprocess")
        self.assertTrue(any(item["name"] == "custom-preprocess" for item in row_domains))

    def test_preprocess_with_input_files_txt_and_docx(self):
        with tempfile.TemporaryDirectory() as tmp:
            txt = os.path.join(tmp, "a.txt")
            docx = os.path.join(tmp, "b.docx")
            out = os.path.join(tmp, "out.jsonl")

            with open(txt, "w", encoding="utf-8") as f:
                f.write("claim one")
            from docx import Document  # type: ignore

            d = Document()
            d.add_paragraph("claim two")
            d.save(docx)

            res = preprocess.preprocess_file(
                txt,
                out,
                {
                    "input_files": [txt, docx],
                    "output_format": "jsonl",
                    "field_transforms": [{"field": "text", "op": "lower"}],
                },
            )
            self.assertEqual(res["summary"]["output_rows"], 2)
            rows = preprocess._read_jsonl(out)
            self.assertEqual(rows[0]["text"], "claim one")

    def test_preprocess_with_input_files_skip_image(self):
        with tempfile.TemporaryDirectory() as tmp:
            txt = os.path.join(tmp, "a.txt")
            img = os.path.join(tmp, "b.png")
            out = os.path.join(tmp, "out.jsonl")
            with open(txt, "w", encoding="utf-8") as f:
                f.write("evidence")
            with open(img, "wb") as f:
                f.write(b"fake")

            res = preprocess.preprocess_file(
                txt,
                out,
                {
                    "input_files": [txt, img],
                    "output_format": "jsonl",
                    "ocr_enabled": False,
                    "on_file_error": "skip",
                },
            )
            self.assertEqual(res["summary"]["output_rows"], 1)
            self.assertTrue(isinstance(res.get("skipped_files"), list))

    def test_preprocess_deduplicate_by(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "cooked.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(json.dumps({"speaker": "alice", "claim_text": "A"}) + "\n")
                f.write(json.dumps({"speaker": "alice", "claim_text": "A"}) + "\n")
                f.write(json.dumps({"speaker": "bob", "claim_text": "B"}) + "\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_format": "jsonl",
                    "output_format": "jsonl",
                    "deduplicate_by": ["speaker", "claim_text"],
                    "deduplicate_keep": "first",
                },
            )
            self.assertEqual(res["summary"]["input_rows"], 3)
            self.assertEqual(res["summary"]["output_rows"], 2)
            self.assertEqual(res["summary"]["duplicate_rows_removed"], 1)

    def test_preprocess_with_input_files_xlsx_all_sheets(self):
        with tempfile.TemporaryDirectory() as tmp:
            xlsx = os.path.join(tmp, "a.xlsx")
            out = os.path.join(tmp, "out.jsonl")
            from openpyxl import Workbook  # type: ignore

            wb = Workbook()
            ws = wb.active
            ws.title = "S1"
            ws.append(["ClaimText"])
            ws.append(["claim one"])
            ws2 = wb.create_sheet("S2")
            ws2.append(["ClaimText"])
            ws2.append(["claim two"])
            wb.save(xlsx)

            res = preprocess.preprocess_file(
                xlsx,
                out,
                {
                    "input_files": [xlsx],
                    "output_format": "jsonl",
                    "xlsx_all_sheets": True,
                    "header_map": {"ClaimText": "claim_text"},
                },
            )
            self.assertEqual(res["summary"]["output_rows"], 2)
            rows = preprocess._read_jsonl(out)
            self.assertEqual({r["sheet_name"] for r in rows}, {"S1", "S2"})

    def test_preprocess_blocks_xlsx_input_when_quality_rules_fail(self):
        with tempfile.TemporaryDirectory() as tmp:
            xlsx = os.path.join(tmp, "a.xlsx")
            out = os.path.join(tmp, "out.jsonl")
            from openpyxl import Workbook  # type: ignore

            wb = Workbook()
            ws = wb.active
            ws.title = "S1"
            ws.append(["id", "note"])
            ws.append([1, "hello"])
            wb.save(xlsx)

            with self.assertRaisesRegex(RuntimeError, "required_columns"):
                preprocess.preprocess_file(
                    xlsx,
                    out,
                    {
                        "input_files": [xlsx],
                        "output_format": "jsonl",
                        "xlsx_rules": {"required_columns": ["id", "amount"]},
                    },
                )

    def test_preprocess_standardize_evidence_and_quality_report(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "cooked.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(json.dumps({"text": "Claim A", "author": "Alice", "source_url": "https://a"}) + "\n")
                f.write(json.dumps({"text": "Claim B", "author": "Bob", "source_url": "https://b"}) + "\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_format": "jsonl",
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "generate_quality_report": True,
                    "quality_required_fields": ["claim_text", "source_url"],
                },
            )
            self.assertTrue(res["quality_report_path"])
            self.assertTrue(os.path.isfile(res["quality_report_path"]))
            rows = preprocess._read_jsonl(dst)
            self.assertIn("evidence_id", rows[0])
            self.assertIn("claim_text", rows[0])
            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertEqual(report["rows"], 2)
            self.assertEqual(report["required_field_missing"]["claim_text"], 0)
            self.assertEqual(report["required_field_missing"]["source_path"], 0)
            self.assertEqual(report["metrics"]["source_ref_coverage"], 1.0)

    def test_preprocess_standardize_evidence_extracts_debate_fields(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "cooked.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(
                    json.dumps(
                        {
                            "text": 'Alice: We should support public transit because it cuts congestion. [1] "Cleaner air"',
                            "source_url": "https://example.com/report",
                        }
                    )
                    + "\n"
                )

            preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_format": "jsonl",
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                },
            )
            rows = preprocess._read_jsonl(dst)
            self.assertEqual(rows[0]["speaker"], "Alice")
            self.assertEqual(rows[0]["speaker_role"], "speaker")
            self.assertEqual(rows[0]["stance"], "pro")
            self.assertEqual(rows[0]["source_domain"], "example.com")
            self.assertEqual(rows[0]["citation_text"], "[1]")
            self.assertEqual(rows[0]["quote_text"], "Cleaner air")
            self.assertEqual(rows[0]["language"], "en")
            self.assertIn("public transit", rows[0]["debate_topic"])

    def test_preprocess_applies_ftfy_repairs_before_cleaning(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(json.dumps({"text": "FranÃ§ais source note"}) + "\n")

            def _fake_fix_text(value, **_kwargs):
                return str(value).replace("FranÃ§ais", "Français")

            with patch("aiwf.preprocess_enrichment._ftfy_fix_text", return_value=_fake_fix_text):
                res = preprocess.preprocess_file(
                    src,
                    dst,
                    {
                        "input_format": "jsonl",
                        "output_format": "jsonl",
                        "generate_quality_report": True,
                        "quality_required_fields": ["text"],
                    },
                )

            rows = preprocess._read_jsonl(dst)
            self.assertEqual(rows[0]["text"], "Français source note")
            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertEqual(report["metrics"]["encoding_repair_ratio"], 1.0)
            self.assertIn("ftfy", report["metrics"]["backend_trace_summary"])

    def test_preprocess_repairs_pdf_radicals_without_ftfy(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(json.dumps({"text": "\u2f50\u8d5b\u2edb\u4e91\x01\u7eff\u2f8a\u7a0e \u2e9f\u516c\u53f8"}) + "\n")

            with patch("aiwf.preprocess_enrichment._ftfy_fix_text", return_value=None):
                res = preprocess.preprocess_file(
                    src,
                    dst,
                    {
                        "input_format": "jsonl",
                        "output_format": "jsonl",
                        "generate_quality_report": True,
                        "quality_required_fields": ["text"],
                    },
                )

            rows = preprocess._read_jsonl(dst)
            self.assertEqual(rows[0]["text"], "比赛风云 绿色税 母公司")
            self.assertEqual(res["summary"]["encoding_rows_repaired"], 1)
            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertEqual(report["metrics"]["encoding_repair_ratio"], 1.0)
            ftfy_summary = report["metrics"]["backend_trace_summary"]["ftfy"]
            self.assertFalse(ftfy_summary["ftfy_available"])
            self.assertTrue(ftfy_summary["builtin_unicode_repair"])

    def test_preprocess_encoding_repair_ratio_uses_input_rows_after_merges(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "wrapped_repair.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("Alice: Low-carbon transition\x01\n")
                f.write("needs coordinated funding plans\x01\n")

            with patch("aiwf.preprocess_enrichment._ftfy_fix_text", return_value=None):
                res = preprocess.preprocess_file(
                    src,
                    dst,
                    {
                        "input_files": [src],
                        "output_format": "jsonl",
                        "standardize_evidence": True,
                        "canonical_profile": "debate_evidence",
                        "generate_quality_report": True,
                        "text_split_by_line": True,
                    },
                )

            self.assertEqual(res["summary"]["input_rows"], 2)
            self.assertEqual(res["summary"]["output_rows"], 1)
            self.assertEqual(res["summary"]["encoding_rows_repaired"], 2)
            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertEqual(report["metrics"]["encoding_repair_ratio"], 1.0)

    def test_preprocess_enriches_url_metadata_only_for_existing_urls(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(json.dumps({"text": "Claim one", "source_url": "https://example.com/report"}) + "\n")
                f.write(json.dumps({"text": "Claim two"}) + "\n")

            def _fake_module_status(name):
                if name == "trafilatura":
                    return True, object()
                return False, None

            with patch("aiwf.preprocess_enrichment._module_status", side_effect=_fake_module_status):
                with patch(
                    "aiwf.preprocess_enrichment._fetch_url_metadata_with_trafilatura",
                    return_value={
                        "ok": True,
                        "title": "Example policy report",
                        "published_at": "2024-09-01",
                        "source_domain": "example.com",
                    },
                ):
                    res = preprocess.preprocess_file(
                        src,
                        dst,
                        {
                            "input_format": "jsonl",
                            "output_format": "jsonl",
                            "standardize_evidence": True,
                            "generate_quality_report": True,
                            "external_enrichment_mode": "private",
                            "url_metadata_enrichment": True,
                        },
                    )

            rows = preprocess._read_jsonl(dst)
            enriched_row = next(row for row in rows if row.get("source_url") == "https://example.com/report")
            self.assertEqual(enriched_row["source_title"], "Example policy report")
            self.assertEqual(enriched_row["published_at"], "2024-09-01")
            self.assertEqual(enriched_row["source_domain"], "example.com")
            self.assertEqual(len([row for row in rows if row.get("source_title") == "Example policy report"]), 1)
            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertEqual(report["metrics"]["url_metadata_resolution_rate"], 1.0)

    def test_preprocess_grobid_fallback_warns_without_blocking(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(
                    json.dumps(
                        {
                            "text": "[1] Example Daily 2024 https://example.com/report",
                            "source_path": os.path.join(tmp, "notes.pdf"),
                        }
                    )
                    + "\n"
                )

            def _fake_module_status(name):
                if name == "grobid_client":
                    return True, object()
                return False, None

            with patch("aiwf.preprocess_enrichment._module_status", side_effect=_fake_module_status):
                with patch("aiwf.preprocess_enrichment._grobid_endpoint", return_value="http://grobid.local"):
                    with patch(
                        "aiwf.preprocess_enrichment._fetch_grobid_citations",
                        return_value={"ok": False, "error": "temporary unavailable"},
                    ):
                        res = preprocess.preprocess_file(
                            src,
                            dst,
                            {
                                "input_format": "jsonl",
                                "output_format": "jsonl",
                                "standardize_evidence": True,
                                "generate_quality_report": True,
                                "external_enrichment_mode": "private",
                                "citation_parse_backend": "grobid",
                            },
                        )

            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertIn("grobid: temporary unavailable", report["warnings"])
            self.assertEqual(report["metrics"]["citation_parse_success_rate"], 1.0)
            self.assertEqual(
                report["metrics"]["backend_trace_summary"]["regex_citation_parser"]["ok"],
                1,
            )
            self.assertEqual(
                report["metrics"]["backend_trace_summary"]["grobid"]["fallbacks"]["regex"],
                1,
            )
            self.assertEqual(
                report["metrics"]["backend_trace_summary"]["regex_citation_parser"]["matched_rows"],
                1,
            )
            self.assertEqual(report["metrics"]["backend_fallback_failures"], [])

    def test_preprocess_grobid_appends_citations_missing_from_local_rows(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "out.jsonl")
            pdf_path = os.path.join(tmp, "notes.pdf")
            with open(src, "w", encoding="utf-8") as f:
                f.write(
                    json.dumps(
                        {
                            "text": "Alice: We support public transit because evidence shows cleaner air.",
                            "source_path": pdf_path,
                            "page": 2,
                        }
                    )
                    + "\n"
                )

            def _fake_module_status(name):
                if name == "grobid_client":
                    return True, object()
                return False, None

            with patch("aiwf.preprocess_enrichment._module_status", side_effect=_fake_module_status):
                with patch("aiwf.preprocess_enrichment._grobid_endpoint", return_value="http://grobid.local"):
                    with patch(
                        "aiwf.preprocess_enrichment._fetch_grobid_citations",
                        return_value={
                            "ok": True,
                            "citations": [
                                "City Mobility Journal 2024 https://city.example/air",
                                "Transit Review 2025 https://transit.example/review",
                            ],
                        },
                    ):
                        res = preprocess.preprocess_file(
                            src,
                            dst,
                            {
                                "input_format": "jsonl",
                                "output_format": "jsonl",
                                "standardize_evidence": True,
                                "generate_quality_report": True,
                                "external_enrichment_mode": "private",
                                "citation_parse_backend": "grobid",
                            },
                        )

            rows = preprocess._read_jsonl(dst)
            citation_rows = [row for row in rows if row.get("argument_role") == "citation"]
            self.assertEqual(len(citation_rows), 2)
            self.assertEqual({row["source_type"] for row in citation_rows}, {"pdf"})
            self.assertTrue(all(row.get("evidence_id") for row in citation_rows))
            self.assertEqual({row.get("page") for row in citation_rows}, {""})
            self.assertEqual(
                {row["source_domain"] for row in citation_rows},
                {"city.example", "transit.example"},
            )
            claim_row = next(row for row in rows if row.get("speaker") == "Alice")
            self.assertEqual(claim_row["page"], 2)
            self.assertEqual(claim_row["source_domain"], "city.example")

            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            grobid_summary = report["metrics"]["backend_trace_summary"]["grobid"]
            self.assertEqual(grobid_summary["ok"], 1)
            self.assertEqual(grobid_summary["matched_rows"], 0)
            self.assertEqual(grobid_summary["appended_rows"], 2)
            self.assertEqual(report["metrics"]["citation_parse_success_rate"], 1.0)

    def test_preprocess_explicit_grobid_blocks_when_regex_fallback_fails(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(
                    json.dumps(
                        {
                            "text": "Alice: We support civic education because it improves trust.",
                            "source_path": os.path.join(tmp, "notes.pdf"),
                        }
                    )
                    + "\n"
                )

            def _fake_module_status(name):
                if name == "grobid_client":
                    return True, object()
                return False, None

            with patch("aiwf.preprocess_enrichment._module_status", side_effect=_fake_module_status):
                with patch("aiwf.preprocess_enrichment._grobid_endpoint", return_value="http://grobid.local"):
                    with patch(
                        "aiwf.preprocess_enrichment._fetch_grobid_citations",
                        return_value={"ok": False, "error": "temporary unavailable"},
                    ):
                        with self.assertRaisesRegex(RuntimeError, "citation backend grobid failed"):
                            preprocess.preprocess_file(
                                src,
                                dst,
                                {
                                    "input_format": "jsonl",
                                    "output_format": "jsonl",
                                    "standardize_evidence": True,
                                    "generate_quality_report": True,
                                    "external_enrichment_mode": "private",
                                    "citation_parse_backend": "grobid",
                                },
                            )

            report_path = os.path.join(tmp, "out.jsonl.quality.json")
            with open(report_path, "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertTrue(report["blocked"])
            self.assertEqual(
                report["metrics"]["backend_fallback_failures"],
                ["citation backend grobid failed and regex fallback parsed no citation rows"],
            )

    def test_preprocess_azure_fallback_warns_without_losing_local_sources(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "out.jsonl")
            pdf_path = os.path.join(tmp, "pack.pdf")
            with open(src, "w", encoding="utf-8") as f:
                f.write(
                    json.dumps(
                        {
                            "text": "Alice: We support civic tax games because they improve recall.",
                            "source_path": pdf_path,
                        }
                    )
                    + "\n"
                )
                f.write(
                    json.dumps(
                        {
                            "text": "Source: Civic Game Lab 2026.1.1 https://example.com/game-lab",
                            "source_path": pdf_path,
                        }
                    )
                    + "\n"
                )

            with patch(
                "aiwf.preprocess_enrichment._fetch_azure_layout",
                return_value={"ok": False, "error": "azure timeout"},
            ):
                res = preprocess.preprocess_file(
                    src,
                    dst,
                    {
                        "input_format": "jsonl",
                        "output_format": "jsonl",
                        "standardize_evidence": True,
                        "generate_quality_report": True,
                        "external_enrichment_mode": "public",
                        "document_parse_backend": "azure_docintelligence",
                    },
                )

            rows = preprocess._read_jsonl(dst)
            claim_row = next(row for row in rows if row.get("speaker") == "Alice")
            self.assertEqual(claim_row["source_title"], "Civic Game Lab")
            self.assertEqual(claim_row["source_url"], "https://example.com/game-lab")
            self.assertEqual(claim_row["published_at"], "2026-01-01")

            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertIn("azure_docintelligence: azure timeout", report["warnings"])
            azure_summary = report["metrics"]["backend_trace_summary"]["azure_docintelligence"]
            self.assertEqual(azure_summary["fallbacks"]["local"], 1)
            self.assertEqual(report["metrics"]["source_ref_coverage"], 1.0)
            self.assertEqual(report["metrics"]["backend_fallback_failures"], [])

    def test_preprocess_explicit_azure_blocks_when_local_fallback_fails(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "out.jsonl")
            pdf_path = os.path.join(tmp, "pack.pdf")
            with open(src, "w", encoding="utf-8") as f:
                f.write(
                    json.dumps(
                        {
                            "text": "Alice: We support civic games because they improve recall.",
                            "source_path": pdf_path,
                            "page": 2,
                        }
                    )
                    + "\n"
                )

            with patch(
                "aiwf.preprocess_enrichment._fetch_azure_layout",
                return_value={"ok": False, "error": "azure timeout"},
            ):
                with self.assertRaisesRegex(RuntimeError, "document backend azure_docintelligence failed"):
                    preprocess.preprocess_file(
                        src,
                        dst,
                        {
                            "input_format": "jsonl",
                            "output_format": "jsonl",
                            "standardize_evidence": True,
                            "generate_quality_report": True,
                            "external_enrichment_mode": "public",
                            "document_parse_backend": "azure_docintelligence",
                        },
                    )

            report_path = os.path.join(tmp, "out.jsonl.quality.json")
            with open(report_path, "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertTrue(report["blocked"])
            self.assertEqual(
                report["metrics"]["backend_fallback_failures"],
                ["document backend azure_docintelligence failed and local fallback produced no source/citation structure"],
            )

    def test_preprocess_azure_layout_roles_promote_structure_rows(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "out.jsonl")
            pdf_path = os.path.join(tmp, "pack.pdf")
            with open(src, "w", encoding="utf-8") as f:
                for text in (
                    "Evidence Pack",
                    "[1] City Mobility Lab 2026 https://city.example/report",
                    "Archive footer",
                ):
                    f.write(json.dumps({"text": text, "source_path": pdf_path}) + "\n")

            azure_payload = {
                "ok": True,
                "payload": {
                    "analyzeResult": {
                        "paragraphs": [
                            {"content": "Evidence Pack", "role": "sectionHeading"},
                            {"content": "[1] City Mobility Lab 2026 https://city.example/report", "role": "footnote"},
                            {"content": "Archive footer", "role": "pageFooter"},
                        ]
                    }
                },
            }
            with patch("aiwf.preprocess_enrichment._fetch_azure_layout", return_value=azure_payload):
                res = preprocess.preprocess_file(
                    src,
                    dst,
                    {
                        "input_format": "jsonl",
                        "output_format": "jsonl",
                        "standardize_evidence": True,
                        "generate_quality_report": True,
                        "external_enrichment_mode": "public",
                        "document_parse_backend": "azure_docintelligence",
                    },
                )

            rows = preprocess._read_jsonl(dst)
            self.assertEqual(next(row for row in rows if row["claim_text"] == "Evidence Pack")["argument_role"], "section")
            citation_row = next(row for row in rows if row["claim_text"] == "City Mobility Lab 2026")
            self.assertEqual(citation_row["argument_role"], "citation")
            self.assertEqual(citation_row["source_title"], "City Mobility Lab 2026")
            self.assertEqual(citation_row["citation_text"], "[1]")
            self.assertEqual(next(row for row in rows if row["claim_text"] == "Archive footer")["argument_role"], "metadata")

            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            azure_summary = report["metrics"]["backend_trace_summary"]["azure_docintelligence"]
            self.assertEqual(azure_summary["ok"], 1)
            self.assertEqual(azure_summary["matched_rows"], 3)

    def test_preprocess_azure_appends_missing_footnote_rows(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "out.jsonl")
            pdf_path = os.path.join(tmp, "pack.pdf")
            with open(src, "w", encoding="utf-8") as f:
                f.write(
                    json.dumps(
                        {
                            "text": "Alice: We support civic games because they improve recall.",
                            "source_path": pdf_path,
                            "page": 2,
                        }
                    )
                    + "\n"
                )

            azure_payload = {
                "ok": True,
                "payload": {
                    "analyzeResult": {
                        "paragraphs": [
                            {"content": "Civic Impact Review 2025.2.1 https://impact.example/review", "role": "footnote"},
                            {"content": "Archived debate footer", "role": "pageFooter"},
                        ]
                    }
                },
            }
            with patch("aiwf.preprocess_enrichment._fetch_azure_layout", return_value=azure_payload):
                res = preprocess.preprocess_file(
                    src,
                    dst,
                    {
                        "input_format": "jsonl",
                        "output_format": "jsonl",
                        "standardize_evidence": True,
                        "generate_quality_report": True,
                        "external_enrichment_mode": "public",
                        "document_parse_backend": "azure_docintelligence",
                    },
                )

            rows = preprocess._read_jsonl(dst)
            citation_row = next(row for row in rows if row.get("argument_role") == "citation")
            self.assertEqual(citation_row["source_title"], "Civic Impact Review")
            self.assertEqual(citation_row["source_url"], "https://impact.example/review")
            self.assertEqual(citation_row["source_type"], "pdf")
            self.assertEqual(citation_row["page"], "")
            metadata_row = next(row for row in rows if row.get("argument_role") == "metadata")
            self.assertEqual(metadata_row["claim_text"], "Archived debate footer")
            claim_row = next(row for row in rows if row.get("speaker") == "Alice")
            self.assertEqual(claim_row["page"], 2)
            self.assertEqual(claim_row["source_domain"], "impact.example")

            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            azure_summary = report["metrics"]["backend_trace_summary"]["azure_docintelligence"]
            self.assertEqual(azure_summary["matched_rows"], 0)
            self.assertEqual(azure_summary["appended_rows"], 2)
            self.assertEqual(report["metrics"]["source_ref_coverage"], 1.0)

    def test_preprocess_citation_rows_do_not_inflate_duplicate_claim_ratio(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(json.dumps({"text": "Alice: We support tax memes because they improve recall."}) + "\n")
                f.write(json.dumps({"text": "[1] Example Daily 2024 https://example.com/report"}) + "\n")
                f.write(json.dumps({"text": "[1] Example Daily 2024 https://example.com/report"}) + "\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_format": "jsonl",
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "generate_quality_report": True,
                    "detect_conflicts": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            citation_rows = [row for row in rows if row.get("argument_role") == "citation"]
            self.assertEqual(len(citation_rows), 1)
            self.assertTrue(all(not row["conflict_flag"] for row in citation_rows))
            self.assertEqual(res["summary"]["structural_duplicate_rows_removed"], 1)
            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertEqual(report["metrics"]["duplicate_claim_ratio"], 0.0)
            self.assertEqual(report["metrics"]["citation_parse_success_rate"], 1.0)

    def test_preprocess_does_not_treat_numbered_claim_items_as_citations(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "numbered_claims.jsonl")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(json.dumps({"text": "1. Tax incentives reduce transition costs."}) + "\n")
                f.write(json.dumps({"text": "(2) Carbon taxes change producer behavior."}) + "\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_format": "jsonl",
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "generate_quality_report": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            self.assertEqual([row.get("argument_role") for row in rows], ["claim", "claim"])
            self.assertTrue(all(not row.get("citation_text") for row in rows))
            self.assertEqual(res["summary"]["citation_candidate_rows"], 0)

    def test_preprocess_keeps_numbered_bibliography_entries_as_citations(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "numbered_sources.jsonl")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(
                    json.dumps({"text": "1. Youth Tax Education Study 2024 https://tax.example/youth"})
                    + "\n"
                )

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_format": "jsonl",
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "generate_quality_report": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            self.assertEqual(len(rows), 1)
            self.assertEqual(rows[0]["argument_role"], "citation")
            self.assertEqual(rows[0]["source_title"], "Youth Tax Education Study 2024")
            self.assertEqual(rows[0]["source_url"], "https://tax.example/youth")
            self.assertEqual(rows[0]["published_at"], "2024-01-01")
            self.assertEqual(res["summary"]["citation_candidate_rows"], 1)

    def test_preprocess_deduplicates_exact_repeated_structural_source_rows_only(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "source_pack.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("Alice: We support civic games because they improve recall.\n")
                f.write("Source: Civic Game Lab 2026.1.1 https://example.com/game-lab\n")
                f.write("Source: Civic Game Lab 2026.1.1 https://example.com/game-lab\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            claim_rows = [row for row in rows if row.get("speaker") == "Alice"]
            citation_rows = [row for row in rows if row.get("argument_role") == "citation"]
            self.assertEqual(len(claim_rows), 1)
            self.assertEqual(len(citation_rows), 1)
            self.assertEqual(claim_rows[0]["source_title"], "Civic Game Lab")
            self.assertEqual(res["summary"]["structural_duplicate_rows_removed"], 1)
            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertEqual(report["summary"]["structural_duplicate_rows_removed"], 1)

    def test_preprocess_source_pack_template_extracts_and_propagates_source_context(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "source_pack.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("【使用意愿】\n")
                f.write("We should keep memes in tax outreach because they improve recall.\n")
                f.write("Young audiences engage more with meme-based tax explainers.\n")
                f.write('——<China Youth Daily: "Young people do not just use memes to stand out">\n')
                f.write("https://example.com/youth\n")

            preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "text_split_by_line": True,
                },
            )
            rows = preprocess._read_jsonl(dst)
            claim_rows = [
                row
                for row in rows
                if row.get("claim_text")
                in {
                    "We should keep memes in tax outreach because they improve recall.",
                    "Young audiences engage more with meme-based tax explainers.",
                }
            ]
            self.assertEqual(len(claim_rows), 2)
            self.assertTrue(
                all(
                    row["source_title"] == 'China Youth Daily: "Young people do not just use memes to stand out"'
                    for row in claim_rows
                )
            )
            self.assertTrue(all(row["source_url"] == "https://example.com/youth" for row in claim_rows))
            self.assertTrue(all(row["source_domain"] == "example.com" for row in claim_rows))

            source_title_row = next(
                row
                for row in rows
                if row.get("claim_text") == 'China Youth Daily: "Young people do not just use memes to stand out"'
            )
            self.assertEqual(source_title_row["argument_role"], "citation")
            self.assertEqual(source_title_row["quote_text"], "")

    def test_preprocess_backfills_source_from_structural_attribution_heading(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "survey_attribution.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("\u3010\u4f7f\u7528\u573a\u5408\u3011\n")
                f.write("\u4e2d\u56fd\u9752\u5e74\u62a5\u793e\u793e\u4f1a\u8c03\u67e5\u4e2d\u5fc3\u8054\u5408\u95ee\u5377\u7f51\u7684\u8c03\u67e5\u663e\u793a:\n")
                f.write("74.6%\u7684\u53d7\u8bbf\u9752\u5e74\u8ba4\u4e3a\u5e94\u5206\u6e05\u573a\u5408\u3001\u7406\u6027\u4f7f\u7528\u7f51\u7edc\u6d41\u884c\u8bed\u3002\n")
                f.write("\u53d7\u8bbf\u9752\u5e74\u8fd8\u5efa\u8bae\u6279\u5224\u63a5\u53d7\u3001\u62b5\u5236\u7c97\u4fd7\u7f51\u7edc\u6897\u3002\n")
                f.write("\u3010\u65b0\u6218\u573a\u3011\n")
                f.write("\u8fd9\u4e00\u884c\u4e0d\u5e94\u7ee7\u627f\u4e0a\u4e00\u4e2a\u6765\u6e90\u3002\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "canonical_profile": "debate_evidence",
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            attributed_rows = [
                row
                for row in rows
                if row.get("claim_text", "").startswith(("74.6%", "\u53d7\u8bbf\u9752\u5e74"))
            ]
            self.assertEqual(len(attributed_rows), 2)
            self.assertTrue(
                all(
                    row["source_title"]
                    == "\u4e2d\u56fd\u9752\u5e74\u62a5\u793e\u793e\u4f1a\u8c03\u67e5\u4e2d\u5fc3\u8054\u5408\u95ee\u5377\u7f51"
                    for row in attributed_rows
                )
            )
            next_section_row = next(row for row in rows if row.get("claim_text") == "\u8fd9\u4e00\u884c\u4e0d\u5e94\u7ee7\u627f\u4e0a\u4e00\u4e2a\u6765\u6e90\u3002")
            self.assertEqual(next_section_row.get("source_title"), "")
            self.assertEqual(res["summary"]["structural_source_context_backfilled_rows"], 2)
            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            source_trace = report["metrics"]["backend_trace_summary"]["local_source_context_backfill"]
            self.assertEqual(source_trace["structural_source_rows"], 2)

    def test_preprocess_backfills_source_across_adjacent_pdf_page_break(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "page_break.jsonl")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(
                    json.dumps(
                        {
                            "text": "Alice: The infrastructure budget increased for environmental facilities.",
                            "page": 1,
                        }
                    )
                    + "\n"
                )
                f.write(
                    json.dumps(
                        {
                            "text": "\u2e3a 2022 CCTV News policy briefing",
                            "source_url": "https://news.example.org/policy-briefing",
                            "page": 2,
                        }
                    )
                    + "\n"
                )

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_format": "jsonl",
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "canonical_profile": "debate_evidence",
                    "generate_quality_report": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            claim_row = next(row for row in rows if row.get("speaker") == "Alice")
            source_row = next(row for row in rows if row.get("argument_role") == "citation")
            self.assertEqual(source_row["source_title"], "CCTV News policy briefing")
            self.assertEqual(claim_row["source_title"], "CCTV News policy briefing")
            self.assertEqual(claim_row["source_url"], "https://news.example.org/policy-briefing")
            self.assertEqual(claim_row["source_domain"], "news.example.org")
            self.assertEqual(res["summary"]["adjacent_page_source_context_backfilled_rows"], 1)
            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            source_trace = report["metrics"]["backend_trace_summary"]["local_source_context_backfill"]
            self.assertEqual(source_trace["adjacent_page_rows"], 1)

    def test_preprocess_collapses_adjacent_citation_title_and_url_rows(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "source_card.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("Usage Intent:\n")
                f.write("Survey shows most respondents use memes willingly.\n")
                f.write("\u2014\u2014<Youth Daily: Meme use is not only novelty>\n")
                f.write("https://example.com/youth\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            claim_row = next(row for row in rows if row.get("argument_role") == "claim")
            self.assertEqual(claim_row["source_title"], "Youth Daily: Meme use is not only novelty")
            self.assertEqual(claim_row["source_url"], "https://example.com/youth")
            citation_rows = [row for row in rows if row.get("argument_role") == "citation"]
            self.assertEqual(len(citation_rows), 1)
            self.assertEqual(citation_rows[0]["source_title"], "Youth Daily: Meme use is not only novelty")
            self.assertEqual(citation_rows[0]["source_url"], "https://example.com/youth")
            self.assertIn("https://example.com/youth", citation_rows[0]["citation_text"])
            self.assertEqual(res["summary"]["adjacent_citation_url_rows_collapsed"], 1)
            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertEqual(report["summary"]["adjacent_citation_url_rows_collapsed"], 1)

    def test_preprocess_derives_source_title_from_url_only_source_cards(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "url_only_source_card.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("The claim needs stable attribution.\n")
                f.write("https://policy.example.org/reports/civic-trust-brief.pdf\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "canonical_profile": "debate_evidence",
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            claim_row = next(row for row in rows if row.get("argument_role") == "claim")
            citation_row = next(row for row in rows if row.get("argument_role") == "citation")
            self.assertEqual(citation_row["source_title"], "Civic Trust Brief")
            self.assertEqual(citation_row["source_url"], "https://policy.example.org/reports/civic-trust-brief.pdf")
            self.assertEqual(claim_row["source_title"], "Civic Trust Brief")
            self.assertEqual(claim_row["source_url"], "https://policy.example.org/reports/civic-trust-brief.pdf")
            self.assertEqual(res["summary"]["source_title_enriched_rows"], 1)

    def test_preprocess_skips_machine_generated_url_slug_for_source_title(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "machine_slug_source_card.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("The policy improved recall among young audiences.\n")
                f.write("https://m.cyol.com/gb/articles/2023-01/11/content_YOOyLBUmpZ.html\n")

            preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "canonical_profile": "debate_evidence",
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            citation_row = next(row for row in rows if row.get("argument_role") == "citation")
            claim_row = next(row for row in rows if row.get("argument_role") == "claim")
            self.assertEqual(citation_row["source_title"], "CYOL")
            self.assertEqual(claim_row["source_title"], "CYOL")

    def test_preprocess_strips_browser_text_fragments_from_source_urls(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "text_fragment_url.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("Young audiences engage more with tax explainers.\n")
                f.write("https://m.cyol.com/gb/articles/2023-01/11/content_YOOyLBUmpZ.html#:~:text=\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "canonical_profile": "debate_evidence",
                    "text_split_by_line": True,
                    "generate_quality_report": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            expected_url = "https://m.cyol.com/gb/articles/2023-01/11/content_YOOyLBUmpZ.html"
            citation_row = next(row for row in rows if row.get("argument_role") == "citation")
            claim_row = next(row for row in rows if row.get("argument_role") == "claim")
            self.assertEqual(citation_row["source_url"], expected_url)
            self.assertEqual(citation_row["citation_text"], expected_url)
            self.assertEqual(claim_row["source_url"], expected_url)
            self.assertNotIn(":~:text", claim_row["citation_text"])
            self.assertEqual(res["summary"]["source_url_normalized_rows"], 1)
            self.assertEqual(res["summary"]["citation_text_url_normalized_rows"], 1)
            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            url_normalizer = report["metrics"]["backend_trace_summary"]["local_url_normalizer"]
            self.assertEqual(url_normalizer["source_url_rows"], 1)
            self.assertEqual(url_normalizer["citation_text_rows"], 1)
            self.assertEqual(report["metrics"]["source_url_normalization_rate"], 1.0)

    def test_preprocess_uses_domain_title_when_url_path_is_machine_id(self):
        cases = [
            ("https://www.nea.gov.cn/2024-07/31/c_1310783380.htm", "NEA"),
            ("https://www.stats.gov.cn/sj/ndsj/2024/indexch.htm", "Stats"),
            ("https://guizhou.chinatax.gov.cn/zfxxgk1/fdzdgknr_5635312/qtxx_5635362/202510/t20251029_8", "Guizhou Chinatax"),
            ("https://www.sohu.com/a/965530101_122301143", "SOHU"),
            ("http://news.cnhubei.com/syzx/p/19361335.html", "Cnhubei"),
            ("https://hbepb.hebei.gov.cn/hbhjt/wap/mtbb/101761621375327.html", "Hbepb Hebei"),
            ("https://fddi.fudan.edu.cn/c7/8d/c21257a771981/page.htm", "FDDI Fudan"),
            ("https://www.solarbe.com/topics/16", "Solarbe"),
            ("https://zhejiang.chinatax.gov.cn/art/2025/7/9/art_13229_639690.html", "Zhejiang Chinatax"),
            ("http://222.76.49.185/npsswj/ssxc/swyx/202511/t20251103_618658.htm", "222.76.49.185"),
        ]
        for url, expected_title in cases:
            with self.subTest(url=url):
                with tempfile.TemporaryDirectory() as tmp:
                    src = os.path.join(tmp, "machine_id_source_card.txt")
                    dst = os.path.join(tmp, "out.jsonl")
                    with open(src, "w", encoding="utf-8") as f:
                        f.write("Photovoltaic poverty alleviation projects generated stable income.\n")
                        f.write(url + "\n")

                    preprocess.preprocess_file(
                        src,
                        dst,
                        {
                            "input_files": [src],
                            "output_format": "jsonl",
                            "standardize_evidence": True,
                            "canonical_profile": "debate_evidence",
                            "text_split_by_line": True,
                        },
                    )

                    rows = preprocess._read_jsonl(dst)
                    citation_row = next(row for row in rows if row.get("argument_role") == "citation")
                    claim_row = next(row for row in rows if row.get("argument_role") == "claim")
                    self.assertEqual(citation_row["source_title"], expected_title)
                    self.assertEqual(claim_row["source_title"], expected_title)

    def test_preprocess_replaces_source_list_marker_claim_with_source_title(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "numbered_source_card.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("1\u3001 https://www.nea.gov.cn/2024-07/31/c_1310783380.htm\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "canonical_profile": "debate_evidence",
                    "text_split_by_line": True,
                    "generate_quality_report": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            self.assertEqual(len(rows), 1)
            self.assertEqual(rows[0]["argument_role"], "citation")
            self.assertEqual(rows[0]["claim_text"], "NEA")
            self.assertEqual(rows[0]["source_title"], "NEA")
            self.assertEqual(rows[0]["citation_text"], "https://www.nea.gov.cn/2024-07/31/c_1310783380.htm")
            self.assertEqual(rows[0]["debate_topic"], "")
            self.assertEqual(res["summary"]["source_marker_claim_replaced_rows"], 1)
            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            marker_summary = report["metrics"]["backend_trace_summary"]["local_source_marker_normalizer"]
            self.assertEqual(marker_summary["source_marker_rows"], 1)

    def test_preprocess_keeps_adjacent_numbered_source_cards_distinct(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "adjacent_numbered_source_cards.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("1\u3001 https://www.nea.gov.cn/2024-07/31/c_1310783380.htm\n")
                f.write("2\u3001 https://paper.people.com.cn/zgnyb/pc/content/202601/19/content_30134486.html\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "canonical_profile": "debate_evidence",
                    "text_split_by_line": True,
                    "generate_quality_report": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            citation_rows = [row for row in rows if row.get("argument_role") == "citation"]
            self.assertEqual(len(citation_rows), 2)
            self.assertEqual({row["claim_text"] for row in citation_rows}, {"NEA", "Paper People"})
            self.assertEqual(res["summary"]["source_marker_claim_replaced_rows"], 2)
            self.assertEqual(res["summary"]["adjacent_citation_url_rows_collapsed"], 0)

    def test_preprocess_appends_multiple_source_cards_to_claim_citation_text(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "multi_source_card_backfill.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("Alice: We support photovoltaic incentives because they reduce rural energy poverty.\n")
                f.write("1\u3001 https://www.nea.gov.cn/2024-07/31/c_1310783380.htm\n")
                f.write("2\u3001 https://paper.people.com.cn/zgnyb/pc/content/202601/19/content_30134486.html\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "canonical_profile": "debate_evidence",
                    "text_split_by_line": True,
                    "generate_quality_report": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            claim_row = next(row for row in rows if row.get("speaker") == "Alice")
            self.assertEqual(claim_row["source_title"], "NEA")
            self.assertEqual(claim_row["source_url"], "https://www.nea.gov.cn/2024-07/31/c_1310783380.htm")
            self.assertIn("https://www.nea.gov.cn/2024-07/31/c_1310783380.htm", claim_row["citation_text"])
            self.assertIn(
                "Paper People https://paper.people.com.cn/zgnyb/pc/content/202601/19/content_30134486.html",
                claim_row["citation_text"],
            )
            self.assertEqual(res["summary"]["source_context_backfilled_rows"], 2)
            self.assertEqual(res["summary"]["multi_source_citation_appended_rows"], 1)
            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertEqual(report["metrics"]["multi_source_citation_appended_rows"], 1)
            source_backfill = report["metrics"]["backend_trace_summary"]["local_source_context_backfill"]
            self.assertEqual(source_backfill["multi_source_citation_rows"], 1)
            self.assertIn(
                "multi-source citation appended rows 1",
                report["raw_signal_hit_summary"]["coverage_notes"],
            )

    def test_preprocess_does_not_treat_structural_labels_as_speakers(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "structural_labels.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("\u653f\u7b56\uff1a\u56fd\u5bb6\u5bf9\u5149\u4f0f\u6276\u8d2b\u9879\u76ee\u7ed9\u4e88\u6240\u5f97\u7a0e\u4f18\u60e0\u3002\n")
                f.write("\u5173\u952e\u7ed3\u8bba\uff1a\u7a0e\u6536\u4f18\u60e0\u53ef\u4ee5\u63d0\u9ad8\u8f6c\u578b\u786e\u5b9a\u6027\u3002\n")
                f.write("Source: Mobility Lab 2024 https://city.example/lab\n")
                f.write("Alice: We support bus lanes.\n")

            preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "canonical_profile": "debate_evidence",
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            structural_row = next(row for row in rows if row.get("claim_text") == "\u56fd\u5bb6\u5bf9\u5149\u4f0f\u6276\u8d2b\u9879\u76ee\u7ed9\u4e88\u6240\u5f97\u7a0e\u4f18\u60e0\u3002")
            conclusion_row = next(row for row in rows if row.get("claim_text") == "\u7a0e\u6536\u4f18\u60e0\u53ef\u4ee5\u63d0\u9ad8\u8f6c\u578b\u786e\u5b9a\u6027\u3002")
            source_row = next(row for row in rows if row.get("argument_role") == "citation")
            speaker_row = next(row for row in rows if row.get("speaker") == "Alice")
            self.assertEqual(structural_row["speaker"], "")
            self.assertEqual(structural_row["speaker_role"], "")
            self.assertEqual(conclusion_row["speaker"], "")
            self.assertEqual(conclusion_row["speaker_role"], "")
            self.assertEqual(source_row["speaker_role"], "source")
            self.assertEqual(speaker_row["speaker_role"], "speaker")

    def test_preprocess_preserves_numbered_source_card_title(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "numbered_source_title_card.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("1. National Energy Agency https://www.nea.gov.cn/2024-07/31/c_1310783380.htm\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "canonical_profile": "debate_evidence",
                    "text_split_by_line": True,
                    "generate_quality_report": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            self.assertEqual(len(rows), 1)
            self.assertEqual(rows[0]["argument_role"], "citation")
            self.assertEqual(rows[0]["claim_text"], "National Energy Agency")
            self.assertEqual(rows[0]["source_title"], "National Energy Agency")
            self.assertEqual(rows[0]["source_url"], "https://www.nea.gov.cn/2024-07/31/c_1310783380.htm")
            self.assertEqual(
                rows[0]["citation_text"],
                "National Energy Agency https://www.nea.gov.cn/2024-07/31/c_1310783380.htm",
            )
            self.assertEqual(rows[0]["debate_topic"], "")
            self.assertEqual(res["summary"]["source_marker_claim_replaced_rows"], 1)
            self.assertEqual(res["summary"]["source_title_enriched_rows"], 1)

    def test_preprocess_backfills_sources_by_unique_citation_token(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "token_footnote.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("Alice: We support bus lanes because commute times fall. [7]\n")
                f.write("References:\n")
                f.write("[7] Civic Mobility Brief 2026 https://city.example/brief\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "canonical_profile": "debate_evidence",
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            claim_row = next(row for row in rows if row.get("speaker") == "Alice")
            self.assertEqual(claim_row["citation_text"], "[7]")
            self.assertEqual(claim_row["source_title"], "Civic Mobility Brief 2026")
            self.assertEqual(claim_row["source_url"], "https://city.example/brief")
            self.assertEqual(res["summary"]["citation_token_source_backfilled_rows"], 1)

    def test_preprocess_extracts_fullwidth_bracket_citation_tokens(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "fullwidth_footnote.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("Alice: We support bus electrification because emissions fall. \uff3bA1\uff3d\n")
                f.write("\uff3bA1\uff3d Mobility Brief 2026 https://city.example/brief\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "canonical_profile": "debate_evidence",
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            claim_row = next(row for row in rows if row.get("speaker") == "Alice")
            citation_row = next(row for row in rows if row.get("argument_role") == "citation")
            self.assertEqual(claim_row["citation_text"], "[A1]")
            self.assertEqual(claim_row["source_title"], "Mobility Brief 2026")
            self.assertEqual(claim_row["source_url"], "https://city.example/brief")
            self.assertEqual(citation_row["citation_text"], "[A1]")
            self.assertEqual(res["summary"]["citation_candidate_rows"], 1)

    def test_preprocess_does_not_treat_chinese_section_brackets_as_citations(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "section_brackets.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("\u3010\u73b0\u72b6\u3011\n")
                f.write("Alice: We support targeted incentives because adoption costs are high.\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "canonical_profile": "debate_evidence",
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            section_row = next(row for row in rows if row.get("claim_text") == "\u3010\u73b0\u72b6\u3011")
            claim_row = next(row for row in rows if row.get("speaker") == "Alice")
            self.assertEqual(section_row["argument_role"], "section")
            self.assertEqual(section_row["citation_text"], "")
            self.assertEqual(claim_row["citation_text"], "")
            self.assertEqual(res["summary"]["citation_candidate_rows"], 0)

    def test_preprocess_skips_ambiguous_citation_token_backfill(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "ambiguous_footnote.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("Alice: We support the proposal because pilots improved uptake. [1]\n")
                f.write("References:\n")
                f.write("[1] City Lab Brief 2026 https://city.example/brief\n")
                f.write("New section:\n")
                f.write("[1] Freight Review 2026 https://freight.example/review\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "canonical_profile": "debate_evidence",
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            claim_row = next(row for row in rows if row.get("speaker") == "Alice")
            self.assertEqual(claim_row["citation_text"], "[1]")
            self.assertEqual(claim_row["source_title"], "")
            self.assertEqual(claim_row["source_url"], "")
            self.assertEqual(res["summary"]["citation_token_source_backfilled_rows"], 0)

    def test_preprocess_backfills_following_claims_from_leading_source_card(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "leading_source_card.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("Source: Civic Game Lab 2026 https://example.com/game-lab\n")
                f.write("Alice: We support civic games because they improve recall.\n")
                f.write("The pilot also increased voluntary filing intent.\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "canonical_profile": "debate_evidence",
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            claim_rows = [row for row in rows if row.get("argument_role") == "claim"]
            self.assertEqual(len(claim_rows), 2)
            self.assertTrue(all(row["source_title"] == "Civic Game Lab 2026" for row in claim_rows))
            self.assertTrue(all(row["source_url"] == "https://example.com/game-lab" for row in claim_rows))
            self.assertEqual(res["summary"]["leading_source_context_backfilled_rows"], 2)
            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            source_trace = report["metrics"]["backend_trace_summary"]["local_source_context_backfill"]
            self.assertEqual(source_trace["matched_rows"], 2)
            self.assertEqual(source_trace["leading_source_rows"], 2)
            self.assertEqual(source_trace["citation_token_rows"], 0)

    def test_preprocess_does_not_forward_fill_numbered_reference_entries(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "numbered_reference_no_forward_fill.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("[1] Civic Game Lab 2026 https://example.com/game-lab\n")
                f.write("Alice: This claim intentionally lacks an inline citation.\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "canonical_profile": "debate_evidence",
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            claim_row = next(row for row in rows if row.get("speaker") == "Alice")
            self.assertEqual(claim_row["source_title"], "")
            self.assertEqual(claim_row["source_url"], "")
            self.assertEqual(res["summary"]["leading_source_context_backfilled_rows"], 0)

    def test_preprocess_does_not_smear_consecutive_citation_sources(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "source_cards.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("Claim one needs the group paper source.\n")
                f.write("\u2014\u2014Group paper\n")
                f.write("\u2014\u2014 2024.9 News Net Economic Daily\n")
                f.write("https://news.example/report\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            group_source = next(row for row in rows if row.get("source_title") == "Group paper")
            news_source = next(row for row in rows if row.get("source_title") == "News Net Economic Daily")
            self.assertEqual(group_source["source_url"], "")
            self.assertEqual(group_source["published_at"], "")
            self.assertEqual(news_source["source_url"], "https://news.example/report")
            self.assertEqual(news_source["published_at"], "2024-09-01")
            self.assertEqual(res["summary"]["adjacent_citation_url_rows_collapsed"], 1)

    def test_preprocess_chinese_source_marker_cleans_title_date_and_url(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "zh_source.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("正方一辩：我们支持短视频税宣，因为能提高触达。\n")
                f.write("来源：新华社税宣观察 2025.3.1 https://news.example.cn/tax\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            claim_row = next(row for row in rows if row.get("speaker") == "正方一辩")
            citation_row = next(row for row in rows if row.get("argument_role") == "citation")
            self.assertEqual(claim_row["source_title"], "新华社税宣观察")
            self.assertEqual(claim_row["source_url"], "https://news.example.cn/tax")
            self.assertEqual(claim_row["source_domain"], "news.example.cn")
            self.assertEqual(claim_row["published_at"], "2025-03-01")
            self.assertEqual(citation_row["source_title"], "新华社税宣观察")
            self.assertIn("https://news.example.cn/tax", citation_row["citation_text"])

            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertEqual(report["metrics"]["citation_parse_success_rate"], 1.0)
            self.assertIn("speaker_prefix", report["raw_signal_hit_summary"]["hit_labels"])
            self.assertIn("stance_marker", report["raw_signal_hit_summary"]["hit_labels"])
            self.assertIn("source_reference", report["raw_signal_hit_summary"]["hit_labels"])
            self.assertIn("debate evidence signals found", report["raw_signal_hit_summary"]["recommendation_reason"])

    def test_preprocess_marks_front_matter_rows_as_metadata(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "front_matter.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("比赛相关信息:\n")
                f.write("时间: 4.11 13:30\n")
                f.write("录制文件: https://meeting.tencent.com/crm/KwXEpvYW84\n")
                f.write("辩题: Carbon tax incentives are more effective.\n")
                f.write("A: We support carbon tax incentives because they reduce transition costs.\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            metadata_rows = [row for row in rows if row.get("argument_role") == "metadata"]
            self.assertEqual(len(metadata_rows), 4)
            self.assertTrue(all((row.get("speaker") or "") == "" for row in metadata_rows))
            self.assertTrue(all(row.get("speaker_role") == "metadata" for row in metadata_rows))
            recording_row = next(row for row in metadata_rows if row.get("source_url"))
            self.assertEqual(recording_row["source_domain"], "meeting.tencent.com")

            claim_row = next(row for row in rows if row.get("argument_role") == "claim")
            self.assertEqual(claim_row["speaker"], "A")
            self.assertEqual(claim_row["stance"], "pro")
            self.assertEqual(claim_row["debate_topic"], "Carbon tax incentives are more effective.")
            self.assertEqual(res["summary"]["metadata_topic_rows_propagated"], 1)

            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertEqual(report["raw_signal_summary"]["speaker_signal_rows"], 1)
            self.assertEqual(report["metrics"]["speaker_coverage"], 1.0)
            self.assertEqual(report["metrics"]["debate_topic_coverage"], 1.0)
            self.assertEqual(report["metrics"]["language_coverage"], 1.0)
            self.assertIn(
                "debate topic coverage 1.000000",
                report["raw_signal_hit_summary"]["coverage_notes"],
            )
            self.assertIn(
                "language coverage 1.000000",
                report["raw_signal_hit_summary"]["coverage_notes"],
            )
            self.assertEqual(report["summary"]["metadata_topic_rows_propagated"], 1)

    def test_preprocess_propagates_front_matter_side_to_unknown_claim_stance(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "front_matter_side.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("\u8fa9\u9898\uff1a\u4f4e\u78b3\u4ea7\u4e1a\u7a0e\u6536\u6fc0\u52b1\u66f4\u6709\u6548\n")
                f.write("\u6301\u65b9\uff1a\u53cd\u65b9\n")
                f.write("\u8c22\u8c22\u4e3b\u5e2d\uff0c\u6211\u65b9\u8ba4\u4e3a\u7a0e\u6536\u6fc0\u52b1\u66f4\u5177\u53ef\u6301\u7eed\u6027\u3002\n")
                f.write("Alice: We support the policy because pilots show uptake.\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            inherited = next(row for row in rows if row.get("claim_text", "").startswith("\u8c22\u8c22\u4e3b\u5e2d"))
            explicit = next(row for row in rows if row.get("speaker") == "Alice")
            self.assertEqual(inherited["stance"], "con")
            self.assertEqual(inherited["debate_topic"], "\u4f4e\u78b3\u4ea7\u4e1a\u7a0e\u6536\u6fc0\u52b1\u66f4\u6709\u6548")
            self.assertEqual(explicit["stance"], "pro")
            self.assertEqual(res["summary"]["metadata_stance_rows_propagated"], 1)
            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertEqual(report["raw_signal_summary"]["stance_signal_rows"], 2)
            self.assertEqual(report["summary"]["metadata_stance_rows_propagated"], 1)

    def test_preprocess_propagates_speech_slot_to_unlabeled_claim_speaker_role(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "speech_slot.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("\u4e00\u8fa9\u7a3f3.0\n")
                f.write("\u8c22\u8c22\u4e3b\u5e2d\uff0c\u6211\u65b9\u8ba4\u4e3a\u7a0e\u6536\u6fc0\u52b1\u66f4\u5177\u53ef\u6301\u7eed\u6027\u3002\n")
                f.write("\u3010\u73b0\u72b6\u3011\n")
                f.write("\u4f4e\u78b3\u4f01\u4e1a\u9700\u8981\u8d44\u91d1\u8f6c\u578b\u3002\n")
                f.write("Alice: We support explicit attribution because it improves traceability.\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            inherited = [
                row
                for row in rows
                if row.get("argument_role") == "claim" and not row.get("speaker")
            ]
            explicit = next(row for row in rows if row.get("speaker") == "Alice")
            self.assertEqual({row["speaker_role"] for row in inherited}, {"first_speaker"})
            self.assertEqual(explicit["speaker_role"], "speaker")
            self.assertEqual(res["summary"]["metadata_speaker_role_rows_propagated"], 2)
            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertEqual(report["summary"]["metadata_speaker_role_rows_propagated"], 2)
            self.assertEqual(report["metrics"]["speaker_coverage"], 0.333333)
            self.assertEqual(report["metrics"]["speaker_role_coverage"], 1.0)
            self.assertEqual(report["metrics"]["speaker_attribution_coverage"], 1.0)
            self.assertIn(
                "speaker attribution coverage 1.000000",
                report["raw_signal_hit_summary"]["coverage_notes"],
            )

    def test_preprocess_propagates_section_argument_role_to_claims(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "argument_role_sections.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("\u3010\u53cd\u9a73\u3011\n")
                f.write("Alice: Their model fails because compliance costs rise.\n")
                f.write("Question: What evidence supports that cost estimate?\n")
                f.write("\u3010\u603b\u7ed3\u9648\u8bcd\u3011\n")
                f.write("Therefore incentives are the more sustainable path.\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            rebuttal_row = next(row for row in rows if row.get("speaker") == "Alice")
            question_row = next(row for row in rows if row.get("argument_role") == "question")
            summary_row = next(row for row in rows if row.get("claim_text") == "Therefore incentives are the more sustainable path.")
            self.assertEqual(rebuttal_row["argument_role"], "rebuttal")
            self.assertEqual(question_row["claim_text"], "What evidence supports that cost estimate?")
            self.assertEqual(summary_row["argument_role"], "summary")
            self.assertEqual(res["summary"]["heading_argument_role_rows_propagated"], 2)
            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertEqual(report["summary"]["heading_argument_role_rows_propagated"], 2)

    def test_preprocess_merges_wrapped_claim_lines_before_source_backfill(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "wrapped_claim.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("Alice: Low-carbon technology has high R&D input and long return cycles\n")
                f.write("so tax incentives lower the cost of transition.\n")
                f.write("Source: Policy Review 2026 https://example.com/policy\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            claim_rows = [row for row in rows if row.get("argument_role") == "claim"]
            self.assertEqual(len(claim_rows), 1)
            self.assertEqual(
                claim_rows[0]["claim_text"],
                "Low-carbon technology has high R&D input and long return cycles so tax incentives lower the cost of transition.",
            )
            self.assertEqual(claim_rows[0]["speaker"], "Alice")
            self.assertEqual(claim_rows[0]["source_url"], "https://example.com/policy")
            self.assertEqual(res["summary"]["wrapped_claim_rows_merged"], 1)
            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertEqual(report["summary"]["wrapped_claim_rows_merged"], 1)

    def test_preprocess_merges_numeric_fact_fragment_wrapped_as_section(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "numeric_fragment.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("2023年,民营企业享受加计扣除金额为2.59万亿元,同比增长12.5%,占全部企业加计扣除金额的\n")
                f.write("75%。\n")
                f.write("⸺2024 光明网\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            claim_rows = [row for row in rows if row.get("argument_role") == "claim"]
            self.assertEqual(len(claim_rows), 1)
            self.assertEqual(
                claim_rows[0]["claim_text"],
                "2023年,民营企业享受加计扣除金额为2.59万亿元,同比增长12.5%,占全部企业加计扣除金额的 75%。",
            )
            self.assertEqual(claim_rows[0]["source_title"], "光明网")
            self.assertEqual(res["summary"]["wrapped_claim_rows_merged"], 1)
            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertEqual(report["summary"]["wrapped_claim_rows_merged"], 1)

    def test_preprocess_does_not_merge_numbered_claim_items(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "numbered_claims.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("Tax incentives reduce transition costs\n")
                f.write("Second, carbon taxes change producer behavior\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            claim_rows = [row for row in rows if row.get("argument_role") == "claim"]
            self.assertEqual(len(claim_rows), 2)
            self.assertEqual(res["summary"]["wrapped_claim_rows_merged"], 0)

    def test_preprocess_marks_outline_rows_as_section(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "outline.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("【现状】\n")
                f.write("2.1 头部企业:\n")
                f.write("A: We support public transit because it reduces congestion.\n")
                f.write("B: We oppose the policy because it raises taxes.\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            section_rows = [row for row in rows if row.get("argument_role") == "section"]
            self.assertEqual([row["claim_text"] for row in section_rows], ["【现状】", "2.1 头部企业:"])
            self.assertTrue(all((row.get("speaker") or "") == "" for row in section_rows))
            self.assertTrue(all(row.get("speaker_role") == "structure" for row in section_rows))

            claim_rows = [row for row in rows if row.get("argument_role") == "claim"]
            self.assertEqual({row["speaker"] for row in claim_rows}, {"A", "B"})
            self.assertEqual({row["stance"] for row in claim_rows}, {"pro", "con"})

            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertEqual(report["metrics"]["speaker_coverage"], 1.0)
            self.assertEqual(report["metrics"]["stance_coverage"], 1.0)

    def test_preprocess_propagates_section_topic_to_following_claim_block(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "section_topic_pack.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("Usage Intent:\n")
                f.write("Alice: We support meme outreach because it improves tax recall.\n")
                f.write("Bob: We oppose meme outreach because it can trivialize services.\n")
                f.write("References:\n")
                f.write("Source: Civic Lab 2026 https://example.com/civic\n")
                f.write("Charlie: We support targeted reminders because they reach young taxpayers.\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                },
            )

            rows = preprocess._read_jsonl(dst)
            claim_by_speaker = {row["speaker"]: row for row in rows if row.get("speaker") in {"Alice", "Bob", "Charlie"}}
            self.assertEqual(claim_by_speaker["Alice"]["debate_topic"], "Usage Intent")
            self.assertEqual(claim_by_speaker["Bob"]["debate_topic"], "Usage Intent")
            self.assertNotEqual(claim_by_speaker["Charlie"]["debate_topic"].lower(), "references")
            self.assertEqual(res["summary"]["section_topic_rows_propagated"], 2)
            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertEqual(report["summary"]["section_topic_rows_propagated"], 2)

    def test_preprocess_blocks_when_speaker_signal_is_lost(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "cooked.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(
                    json.dumps(
                        {
                            "Speaker": "Alice",
                            "ClaimText": "Tax policy should be expanded.",
                            "SourceUrl": "https://example.com/report",
                        }
                    )
                    + "\n"
                )

            with self.assertRaisesRegex(RuntimeError, "speaker is required when source rows contain speaker signal"):
                preprocess.preprocess_file(
                    src,
                    dst,
                    {
                        "input_format": "jsonl",
                        "output_format": "jsonl",
                        "standardize_evidence": True,
                        "header_map": {
                            "Speaker": "speaker",
                            "ClaimText": "claim_text",
                            "SourceUrl": "source_url",
                        },
                        "include_fields": ["claim_text", "source_url", "source_path"],
                    },
                )

    def test_preprocess_blocks_when_source_signal_is_lost(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "cooked.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(
                    json.dumps(
                        {
                            "ClaimText": "Tax policy should be expanded.",
                            "SourceTitle": "Policy Review",
                            "SourceUrl": "https://example.com/report",
                        }
                    )
                    + "\n"
                )

            with self.assertRaisesRegex(
                RuntimeError,
                "source_url or source_title is required when source rows contain source reference signal",
            ):
                preprocess.preprocess_file(
                    src,
                    dst,
                    {
                        "input_format": "jsonl",
                        "output_format": "jsonl",
                        "standardize_evidence": True,
                        "header_map": {
                            "ClaimText": "claim_text",
                            "SourceTitle": "source_title",
                            "SourceUrl": "source_url",
                        },
                        "include_fields": ["claim_text", "source_path"],
                    },
                )

    def test_preprocess_debate_deduplicate_keeps_distinct_speaker_or_source_rows(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "cooked.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(
                    json.dumps(
                        {
                            "text": "Alice: Bus lanes cut congestion.",
                            "source_url": "https://example.com/report-a",
                        }
                    )
                    + "\n"
                )
                f.write(
                    json.dumps(
                        {
                            "text": "Bob: Bus lanes cut congestion.",
                            "source_url": "https://example.com/report-a",
                        }
                    )
                    + "\n"
                )
                f.write(
                    json.dumps(
                        {
                            "text": "Alice: Bus lanes cut congestion.",
                            "source_url": "https://example.org/report-b",
                        }
                    )
                    + "\n"
                )
                f.write(
                    json.dumps(
                        {
                            "text": "Alice: Bus lanes cut congestion.",
                            "source_url": "https://example.com/report-a",
                        }
                    )
                    + "\n"
                )

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_format": "jsonl",
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "deduplicate_by": ["source_path", "speaker", "source_domain", "claim_text"],
                    "deduplicate_keep": "first",
                },
            )

            self.assertEqual(res["summary"]["input_rows"], 4)
            self.assertEqual(res["summary"]["output_rows"], 3)
            self.assertEqual(res["summary"]["duplicate_rows_removed"], 1)
            rows = preprocess._read_jsonl(dst)
            self.assertEqual(
                {
                    (row["speaker"], row["source_domain"], row["claim_text"])
                    for row in rows
                },
                {
                    ("Alice", "example.com", "Bus lanes cut congestion."),
                    ("Bob", "example.com", "Bus lanes cut congestion."),
                    ("Alice", "example.org", "Bus lanes cut congestion."),
                },
            )
            self.assertEqual(len({row["evidence_id"] for row in rows}), 3)

    def test_preprocess_text_split_by_line_strips_ocr_noise_and_preserves_source_url(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "ocr.txt")
            dst = os.path.join(tmp, "out.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write("Page 3\n")
                f.write("Speaker A: We support public transit because it reduces congggggestion.\n")
                f.write("Speaker A: We support public transit because it reduces congggggestion.\n")
                f.write("Source: https://example.com/report\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_files": [src],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "generate_quality_report": True,
                    "text_split_by_line": True,
                    "field_transforms": [{"field": "text", "op": "strip_ocr_noise"}],
                    "row_filters": [{"field": "text", "op": "exists"}],
                    "deduplicate_by": ["source_path", "text"],
                },
            )

            self.assertEqual(res["summary"]["input_rows"], 4)
            self.assertEqual(res["summary"]["output_rows"], 2)
            self.assertEqual(res["summary"]["dropped_by_filters"], 1)
            self.assertEqual(res["summary"]["duplicate_rows_removed"], 1)
            rows = preprocess._read_jsonl(dst)
            self.assertEqual(rows[0]["claim_text"], "We support public transit because it reduces congestion.")
            self.assertEqual(rows[0]["speaker"], "A")
            self.assertEqual(rows[1]["source_url"], "https://example.com/report")
            self.assertEqual(rows[1]["claim_text"], "https://example.com/report")

            with open(res["quality_report_path"], "r", encoding="utf-8") as f:
                report = json.load(f)
            self.assertEqual(report["metrics"]["duplicate_claim_ratio"], 0.0)
            self.assertEqual(report["warnings"], [])

    def test_preprocess_mixed_input_files_share_debate_schema(self):
        with tempfile.TemporaryDirectory() as tmp:
            txt = os.path.join(tmp, "a.txt")
            docx = os.path.join(tmp, "b.docx")
            out = os.path.join(tmp, "out.jsonl")
            with open(txt, "w", encoding="utf-8") as f:
                f.write("正方一辩：应该支持公共交通，因为可以缓解拥堵。")
            from docx import Document  # type: ignore

            d = Document()
            d.add_paragraph('Moderator: "Cleaner air" supports the affirmative case.')
            d.save(docx)

            preprocess.preprocess_file(
                txt,
                out,
                {
                    "input_files": [txt, docx],
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                },
            )
            rows = preprocess._read_jsonl(out)
            self.assertEqual({row["source_type"] for row in rows}, {"txt", "docx"})
            self.assertTrue(all("claim_text" in row for row in rows))
            self.assertTrue(all("source_path" in row for row in rows))
            self.assertTrue(any(row.get("speaker") == "正方一辩" for row in rows))
            self.assertTrue(any(row.get("argument_role") == "quote" for row in rows))

    def test_preprocess_export_canonical_bundle(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "cooked.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(json.dumps({"text": "论文主体第一段", "source_file": "a.pdf"}) + "\n")
                f.write(json.dumps({"text": "论文主体第二段", "source_file": "a.pdf"}) + "\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_format": "jsonl",
                    "output_format": "jsonl",
                    "export_canonical_bundle": True,
                    "canonical_title": "论文熟肉包",
                },
            )
            bundle = res.get("canonical_bundle") or {}
            self.assertTrue(bundle.get("bundle_dir"))
            self.assertTrue(os.path.isfile(bundle.get("markdown_path")))
            self.assertTrue(os.path.isfile(bundle.get("metadata_path")))
            self.assertTrue(os.path.isfile(bundle.get("lineage_path")))
            with open(bundle.get("metadata_path"), "r", encoding="utf-8") as f:
                meta = json.load(f)
            self.assertEqual(meta.get("row_count"), 2)

    def test_preprocess_rejects_relative_canonical_bundle_path_escape(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "cooked.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(json.dumps({"text": "hello", "source_file": "a.txt"}) + "\n")

            with self.assertRaises(ValueError):
                preprocess.preprocess_file(
                    src,
                    dst,
                    {
                        "input_format": "jsonl",
                        "output_format": "jsonl",
                        "export_canonical_bundle": True,
                        "canonical_bundle_dir": r"..\bundle_outside",
                    },
                )

    def test_preprocess_rejects_absolute_quality_report_path_escape(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "cooked.jsonl")
            outside = os.path.join(tempfile.gettempdir(), "aiwf_quality_escape.json")
            with open(src, "w", encoding="utf-8") as f:
                f.write(json.dumps({"text": "hello", "source_file": "a.txt"}) + "\n")

            with self.assertRaises(ValueError):
                preprocess.preprocess_file(
                    src,
                    dst,
                    {
                        "input_format": "jsonl",
                        "output_format": "jsonl",
                        "generate_quality_report": True,
                        "quality_report_path": outside,
                    },
                )

    def test_preprocess_rejects_absolute_canonical_bundle_path_escape(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "cooked.jsonl")
            outside = os.path.join(tempfile.gettempdir(), "aiwf_bundle_escape")
            with open(src, "w", encoding="utf-8") as f:
                f.write(json.dumps({"text": "hello", "source_file": "a.txt"}) + "\n")

            with self.assertRaises(ValueError):
                preprocess.preprocess_file(
                    src,
                    dst,
                    {
                        "input_format": "jsonl",
                        "output_format": "jsonl",
                        "export_canonical_bundle": True,
                        "canonical_bundle_dir": outside,
                    },
                )

    def test_preprocess_chunk_and_conflict_detection(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "cooked.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(json.dumps({"text": "Tax policy is good. We should support it."}) + "\n")
                f.write(json.dumps({"text": "Tax policy is bad. We should oppose it."}) + "\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_format": "jsonl",
                    "output_format": "jsonl",
                    "header_map": {"text": "claim_text"},
                    "chunk_mode": "sentence",
                    "chunk_field": "claim_text",
                    "detect_conflicts": True,
                    "conflict_text_field": "claim_text",
                    "conflict_positive_words": ["support", "good"],
                    "conflict_negative_words": ["oppose", "bad"],
                },
            )
            self.assertGreaterEqual(res["summary"]["chunked_rows_created"], 2)
            self.assertGreaterEqual(res["summary"]["conflict_rows_marked"], 2)
            rows = preprocess._read_jsonl(dst)
            self.assertTrue(any(r.get("conflict_flag") for r in rows))

    def test_preprocess_conflict_detection_groups_same_topic_with_different_reasons(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "cooked.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(json.dumps({"text": "Alice: We support bus lanes because they reduce commute times."}) + "\n")
                f.write(json.dumps({"text": "Bob: We oppose bus lanes because they hurt curbside deliveries."}) + "\n")

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_format": "jsonl",
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "detect_conflicts": True,
                    "conflict_text_field": "claim_text",
                },
            )
            self.assertEqual(res["summary"]["conflict_rows_marked"], 2)
            rows = preprocess._read_jsonl(dst)
            flagged = [row for row in rows if row.get("conflict_flag")]
            self.assertEqual(len(flagged), 2)
            self.assertEqual({row.get("conflict_polarity") for row in flagged}, {"pro", "con"})
            self.assertEqual({row.get("conflict_topic") for row in flagged}, {"bus lanes"})

    def test_preprocess_conflict_detection_skips_quote_only_rows(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "cooked.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(json.dumps({"text": "Alice: We support bus lanes because they reduce commute times."}) + "\n")
                f.write(json.dumps({"text": "Bob: We oppose bus lanes because they hurt curbside deliveries."}) + "\n")
                f.write(json.dumps({"text": 'Reporter: "Bus lanes" was the phrase both teams repeated.'}) + "\n")

            preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_format": "jsonl",
                    "output_format": "jsonl",
                    "standardize_evidence": True,
                    "detect_conflicts": True,
                    "conflict_text_field": "claim_text",
                },
            )
            rows = preprocess._read_jsonl(dst)
            quote_row = next(row for row in rows if row.get("argument_role") == "quote")
            self.assertFalse(quote_row.get("conflict_flag"))
            self.assertEqual(quote_row.get("conflict_polarity"), "unknown")
            flagged = [row for row in rows if row.get("conflict_flag")]
            self.assertEqual(len(flagged), 2)

    def test_preprocess_conflict_detection_skips_citation_rows(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "raw.jsonl")
            dst = os.path.join(tmp, "cooked.jsonl")
            with open(src, "w", encoding="utf-8") as f:
                f.write(
                    json.dumps(
                        {
                            "claim_text": "[1] City Lab supports bus lanes.",
                            "argument_role": "citation",
                            "stance": "pro",
                            "debate_topic": "bus lanes",
                        }
                    )
                    + "\n"
                )
                f.write(
                    json.dumps(
                        {
                            "claim_text": "Bob opposes bus lanes because they hurt deliveries.",
                            "argument_role": "claim",
                            "stance": "con",
                            "debate_topic": "bus lanes",
                        }
                    )
                    + "\n"
                )

            res = preprocess.preprocess_file(
                src,
                dst,
                {
                    "input_format": "jsonl",
                    "output_format": "jsonl",
                    "detect_conflicts": True,
                    "conflict_text_field": "claim_text",
                },
            )
            self.assertEqual(res["summary"]["conflict_rows_marked"], 0)
            rows = preprocess._read_jsonl(dst)
            self.assertTrue(all(not row.get("conflict_flag") for row in rows))
            citation_row = next(row for row in rows if row.get("argument_role") == "citation")
            self.assertEqual(citation_row.get("conflict_polarity"), "pro")

    def test_preprocess_pipeline_extract_clean_structure_audit(self):
        with tempfile.TemporaryDirectory() as tmp:
            txt = os.path.join(tmp, "raw.txt")
            with open(txt, "w", encoding="utf-8") as f:
                f.write("Tax policy should be supported.\nTax policy should be opposed.")

            out_csv = os.path.join(tmp, "final.csv")
            res = preprocess.run_preprocess_pipeline(
                pipeline={
                    "stages": [
                        {"name": "extract", "config": {"input_files": [txt]}},
                        {"name": "clean", "config": {"field_transforms": [{"field": "text", "op": "trim"}]}},
                        {
                            "name": "structure",
                            "config": {
                                "header_map": {"text": "claim_text"},
                                "chunk_mode": "sentence",
                                "chunk_field": "claim_text",
                            },
                        },
                        {
                            "name": "audit",
                            "config": {
                                "detect_conflicts": True,
                                "conflict_text_field": "claim_text",
                                "conflict_positive_words": ["support"],
                                "conflict_negative_words": ["oppose"],
                            },
                        },
                    ]
                },
                job_root=tmp,
                stage_dir=os.path.join(tmp, "stage"),
                input_path=txt,
                final_output_path=out_csv,
            )
            self.assertEqual(res["mode"], "pipeline")
            self.assertEqual(len(res["stages"]), 4)
            self.assertTrue(os.path.isfile(out_csv))
            rows, _ = preprocess._read_csv(out_csv)
            self.assertGreaterEqual(len(rows), 2)

    def test_validate_preprocess_pipeline_rejects_unknown_stage(self):
        vr = preprocess.validate_preprocess_pipeline({"stages": [{"name": "missing_stage", "config": {}}]})
        self.assertFalse(vr["ok"])

    def test_write_rows_supports_bare_filename_outputs(self):
        with tempfile.TemporaryDirectory() as tmp:
            cwd = os.getcwd()
            os.chdir(tmp)
            try:
                out_format = preprocess._write_rows("rows.jsonl", [{"text": "hello"}], {})
                rows = preprocess._read_jsonl("rows.jsonl")
            finally:
                os.chdir(cwd)

        self.assertEqual(out_format, "jsonl")
        self.assertEqual(rows[0]["text"], "hello")

    def test_preprocess_pipeline_supports_registered_custom_stage(self):
        def prepare_uppercase_stage(context):
            cfg = dict(context.config)
            cfg.setdefault("output_format", "jsonl")
            transforms = list(cfg.get("field_transforms") or [])
            transforms.append({"field": "text", "op": "upper"})
            cfg["field_transforms"] = transforms
            return cfg

        preprocess.register_pipeline_stage("uppercase_stage", prepare_config=prepare_uppercase_stage)
        try:
            with tempfile.TemporaryDirectory() as tmp:
                txt = os.path.join(tmp, "raw.txt")
                with open(txt, "w", encoding="utf-8") as f:
                    f.write("hello world")

                pipeline = {
                    "stages": [
                        {"name": "extract", "config": {"input_files": [txt]}},
                        {"name": "uppercase_stage", "config": {}},
                    ]
                }
                vr = preprocess.validate_preprocess_pipeline(pipeline)
                self.assertTrue(vr["ok"])

                out_csv = os.path.join(tmp, "final.csv")
                res = preprocess.run_preprocess_pipeline(
                    pipeline=pipeline,
                    job_root=tmp,
                    stage_dir=os.path.join(tmp, "stage"),
                    input_path=txt,
                    final_output_path=out_csv,
                )
                rows, _ = preprocess._read_csv(out_csv)
        finally:
            preprocess.unregister_pipeline_stage("uppercase_stage")

        self.assertEqual(res["stages"][1]["stage"], "uppercase_stage")
        self.assertEqual(rows[0]["text"], "HELLO WORLD")

    def test_preprocess_pipeline_final_output_respects_requested_format(self):
        with tempfile.TemporaryDirectory() as tmp:
            txt = os.path.join(tmp, "raw.txt")
            with open(txt, "w", encoding="utf-8") as f:
                f.write("hello world")

            out_jsonl = os.path.join(tmp, "final.jsonl")
            res = preprocess.run_preprocess_pipeline(
                pipeline={"stages": [{"name": "extract", "config": {"input_files": [txt]}}]},
                job_root=tmp,
                stage_dir=os.path.join(tmp, "stage"),
                input_path=txt,
                final_output_path=out_jsonl,
            )
            rows = preprocess._read_jsonl(out_jsonl)

        self.assertEqual(res["output_path"], out_jsonl)
        self.assertEqual(rows[0]["text"], "hello world")

    def test_preprocess_pipeline_rejects_final_output_path_escape(self):
        with tempfile.TemporaryDirectory() as tmp:
            txt = os.path.join(tmp, "raw.txt")
            with open(txt, "w", encoding="utf-8") as f:
                f.write("hello world")

            with self.assertRaises(ValueError):
                preprocess.run_preprocess_pipeline(
                    pipeline={"stages": [{"name": "extract", "config": {"input_files": [txt]}}]},
                    job_root=tmp,
                    stage_dir=os.path.join(tmp, "stage"),
                    input_path=txt,
                    final_output_path=r"..\outside.csv",
                )


if __name__ == "__main__":
    unittest.main()
